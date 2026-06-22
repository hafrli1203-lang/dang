# -*- coding: utf-8 -*-
"""AI 상담/브리핑 순수 로직 테스트(네트워크/AI 없이)."""
import io
import unittest

from app.research import briefing as B


class TestExtractText(unittest.TestCase):
    def test_txt_utf8(self):
        data = "변색렌즈 0원 행사\n10월 한 달".encode("utf-8")
        self.assertIn("변색렌즈", B.extract_text("event.txt", data))

    def test_txt_cp949(self):
        data = "다초점렌즈 행사".encode("cp949")
        self.assertEqual(B.extract_text("a.txt", data), "다초점렌즈 행사")

    def test_docx(self):
        from docx import Document
        doc = Document()
        doc.add_paragraph("지니스안경 변색렌즈 행사")
        doc.add_paragraph("10월 한정 0원")
        buf = io.BytesIO()
        doc.save(buf)
        out = B.extract_text("brief.docx", buf.getvalue())
        self.assertIn("변색렌즈", out)
        self.assertIn("0원", out)

    def test_unsupported_raises(self):
        with self.assertRaises(B.BriefingUnsupported):
            B.extract_text("poster.hwp", b"x")

    def test_image_routes_through_cli_provider(self):
        # 이미지는 구독 CLI 멀티모달(image 인자)로 — API 키 불필요.
        import app.ai.providers as P
        orig = P.get_provider
        captured = {}

        class _Fake:
            def generate_text(self, prompt, *, system_prompt=None, image=None, image_mime="image/png"):
                captured["image"] = image
                captured["mime"] = image_mime
                captured["prompt"] = prompt
                return "행사: 7-9월 프로모션\n호야 퓨어젠 할인"
        P.get_provider = lambda engine="claude": _Fake()
        try:
            out = B.extract_text("poster.png", b"PNGDATA")
        finally:
            P.get_provider = orig
        self.assertEqual(captured["image"], b"PNGDATA")
        self.assertEqual(captured["mime"], "image/png")
        self.assertIn("프로모션", out)

    def test_image_unsupported_when_no_image_arg(self):
        # image 인자를 안 받는 백엔드면 친절히 안내(BriefingUnsupported).
        import app.ai.providers as P
        orig = P.get_provider

        class _NoImg:
            def generate_text(self, prompt, *, system_prompt=None):
                return "x"
        P.get_provider = lambda engine="claude": _NoImg()
        try:
            with self.assertRaises(B.BriefingUnsupported):
                B.extract_text("poster.jpg", b"x")
        finally:
            P.get_provider = orig


class TestChatPrompt(unittest.TestCase):
    def test_includes_brief_history_and_msg(self):
        p = B.build_chat_prompt(
            "변색렌즈 0원 행사",
            [("user", "안녕"), ("ai", "네 안녕하세요")],
            "광고로 어떻게 풀지?",
        )
        self.assertIn("변색렌즈 0원 행사", p)
        self.assertIn("사용자: 안녕", p)
        self.assertIn("상담가: 네 안녕하세요", p)
        self.assertIn("광고로 어떻게 풀지?", p)

    def test_default_question_when_empty_msg(self):
        p = B.build_chat_prompt("내용", [], "")
        self.assertIn("광고 관점에서 정리", p)


class TestKeywords(unittest.TestCase):
    def test_parse_strips_bullets_and_dedup(self):
        text = "- 변색렌즈\n2) 다초점렌즈\n변색렌즈\n* 안경 렌즈 추천\n키워드:"
        kws = B.parse_keywords(text)
        self.assertEqual(kws, ["변색렌즈", "다초점렌즈", "안경 렌즈 추천"])

    def test_parse_caps_six(self):
        text = "\n".join(f"소재{i}단어" for i in range(10))
        self.assertEqual(len(B.parse_keywords(text)), 6)

    def test_extract_uses_injected_ai(self):
        captured = {}

        def fake_gen(prompt, system_prompt=None):
            captured["prompt"] = prompt
            return "변색렌즈\n다초점렌즈"
        kws = B.extract_keywords("변색렌즈 0원 행사 공주점", fake_gen)
        self.assertEqual(kws, ["변색렌즈", "다초점렌즈"])
        self.assertIn("변색렌즈 0원 행사", captured["prompt"])

    def test_extract_empty_brief(self):
        self.assertEqual(B.extract_keywords("", lambda *a, **k: "x"), [])

    def test_extract_degrades_on_ai_error(self):
        def boom(*a, **k):
            raise RuntimeError("ai down")
        self.assertEqual(B.extract_keywords("내용", boom), [])


if __name__ == "__main__":
    unittest.main()
