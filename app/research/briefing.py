# -*- coding: utf-8 -*-
"""AI 상담/브리핑 — 순수 로직(텍스트 추출 + 프롬프트 빌드 + 키워드 추출).

행사·캠페인 내용을 붙여넣거나 파일(.txt/.docx, pypdf 있으면 .pdf)로 받아 텍스트로 만들고,
Claude Opus가 읽고 대화·분석한다. AI 호출은 generate_text 콜러블로 주입 → 네트워크 없이 테스트.
"""
from __future__ import annotations

import io
import re

MAX_BRIEF_CHARS = 40_000      # AI에 넣는 브리핑 본문 상한(토큰 보호)
MAX_HISTORY_TURNS = 12        # 프롬프트에 싣는 최근 대화 턴 수


class BriefingUnsupported(Exception):
    """지원하지 않는 파일 형식(또는 라이브러리 미설치)."""


SYSTEM_GUIDE_BRIEFING = """\
당신은 한국 동네 자영업 당근(당근마켓) 광고를 돕는 실전 기획 상담가입니다.
사용자가 올린 '행사/캠페인 내용'을 읽고, 광고로 풀 수 있게 같이 정리합니다.

[원칙]
- 추상어(혁신/차별화/고객가치) 금지. 손님이 실제로 쓰는 구체적인 말로.
- 모르면 모른다고 하고, 더 필요한 정보는 콕 집어 되묻습니다.
- 행사 내용에서 '광고 소재(상품/혜택)'와 '타겟 손님'을 분명히 잡아줍니다.
- 길게 늘어놓지 말고, 핵심부터 짧게. 표/불릿으로 정리해도 좋습니다.
"""


_IMAGE_MIME = {
    ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
    ".webp": "image/webp", ".gif": "image/gif",
}


def extract_text(filename: str, data: bytes) -> str:
    """업로드 파일 바이트 → 텍스트. .txt/.docx, pypdf 있으면 .pdf, 이미지는 Claude 비전. 그 외 Unsupported."""
    name = (filename or "").lower().strip()
    if name.endswith(".txt") or name.endswith(".md") or name.endswith(".csv"):
        return _decode_text(data)
    if name.endswith(".docx"):
        return _extract_docx(data)
    if name.endswith(".pdf"):
        return _extract_pdf(data)
    for ext, mime in _IMAGE_MIME.items():
        if name.endswith(ext):
            return extract_image_text(data, mime)
    raise BriefingUnsupported(
        f"'{filename}'는 아직 지원 안 해요. 텍스트(.txt)·워드(.docx)·PDF·이미지로 올리거나 내용을 붙여넣어 주세요."
    )


_IMAGE_READ_PROMPT = (
    "이 이미지(행사 포스터/전단/안내문 등)에서 '손님에게 광고로 보여줄' 내용만 "
    "한국어 텍스트로 옮겨줘.\n"
    "- 포함: 행사명·기간, 상품/혜택/할인·경품, 대상, 매장/브랜드, 손님 안내 문구.\n"
    "- 제외(절대 옮기지 마): 발주량·매입가·원가·마진·수당·정산·내부 코드 등 "
    "사내 운영/원가 정보는 광고에 못 쓰니 빼라.\n"
    "- 네 설명·해석은 빼고, 적힌 내용 그대로."
)


def extract_image_text(data: bytes, mime: str = "image/png") -> str:
    """이미지(포스터·전단)를 Claude로 읽어 '광고에 쓸 행사·혜택'만 텍스트로 옮긴다.

    구독 CLI 멀티모달(get_provider('claude').generate_text의 image 인자)을 쓴다 —
    API 키·크레딧 불필요(텍스트 채팅과 같은 구독 경로). image 인자를 지원 않는
    백엔드면 BriefingUnsupported.
    """
    from app.ai.providers import get_provider
    try:
        text = get_provider("claude").generate_text(
            _IMAGE_READ_PROMPT, image=data, image_mime=mime,
        )
    except TypeError as e:  # provider가 image 인자를 지원하지 않음(예: API 백엔드)
        raise BriefingUnsupported(
            "현재 AI 백엔드가 이미지 읽기를 지원하지 않아요. 내용을 붙여넣어 주세요."
        ) from e
    return (text or "").strip()


def _decode_text(data: bytes) -> str:
    for enc in ("utf-8-sig", "utf-8", "cp949", "euc-kr"):
        try:
            return data.decode(enc)
        except (UnicodeDecodeError, AttributeError):
            continue
    return data.decode("utf-8", errors="replace")


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:  # 앱엔 python-docx가 있으나 방어적으로
        raise BriefingUnsupported("워드(.docx) 처리 모듈이 없어요.") from e
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def _extract_pdf(data: bytes) -> str:
    try:
        import pypdf
    except ImportError as e:
        raise BriefingUnsupported(
            "PDF는 'pypdf' 설치가 필요해요. 'pip install pypdf' 후 다시 시도하거나, "
            "내용을 붙여넣어 주세요."
        ) from e
    reader = pypdf.PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages).strip()


def build_chat_prompt(brief: str, history: list, user_msg: str) -> str:
    """브리핑 본문 + 최근 대화 + 사용자 질문을 하나의 프롬프트로.

    history: [(role, text)] — role은 'user' | 'ai'.
    """
    brief = (brief or "").strip()[:MAX_BRIEF_CHARS]
    lines: list[str] = []
    if brief:
        lines.append("[행사/캠페인 브리핑 — 사용자가 올린 내용]")
        lines.append(brief)
        lines.append("")
    recent = [t for t in (history or []) if t and len(t) == 2][-MAX_HISTORY_TURNS:]
    if recent:
        lines.append("[지금까지 대화]")
        for role, text in recent:
            who = "사용자" if role == "user" else "상담가"
            lines.append(f"{who}: {str(text).strip()}")
        lines.append("")
    lines.append("[사용자 질문]")
    lines.append((user_msg or "").strip() or "이 내용을 광고 관점에서 정리해 줘.")
    return "\n".join(lines)


_KW_CLEAN = re.compile(r"^[\s\-*•\d.)\]]+")


def extract_keywords(brief: str, generate_text) -> list[str]:
    """브리핑에서 커뮤니티 리서치용 '소재/상품 키워드'를 AI로 뽑는다(지역·금액 제외).

    generate_text(prompt, system_prompt=...) 주입. 실패/빈 결과면 [].
    """
    brief = (brief or "").strip()[:MAX_BRIEF_CHARS]
    if not brief:
        return []
    prompt = (
        "다음 행사/캠페인 내용에서, 손님이 네이버·커뮤니티에서 실제로 검색하거나 이야기할 "
        "'소재/상품 키워드'만 3~6개 뽑아라.\n"
        "- 지역명, 매장 고유명, 금액·할인어(0원/50%/무료 등)는 제외한다.\n"
        "- 손님 입장의 상품/카테고리/고민 단어로(예: 변색렌즈, 다초점렌즈, 안경 렌즈 추천).\n"
        "- 다른 설명 없이 한 줄에 하나씩만 출력한다.\n\n"
        f"[내용]\n{brief}"
    )
    try:
        out = generate_text(prompt, system_prompt=SYSTEM_GUIDE_BRIEFING) or ""
    except Exception:  # noqa: BLE001 — 키워드 추출 실패해도 상담은 정상
        return []
    return parse_keywords(out)


def parse_keywords(text: str) -> list[str]:
    """AI 출력 텍스트 → 키워드 리스트(불릿·번호 제거, 중복제거, 최대 6개)."""
    kws: list[str] = []
    for raw in (text or "").splitlines():
        t = _KW_CLEAN.sub("", raw).strip().strip('"').strip("'").strip()
        if len(t) < 2 or len(t) > 30 or t in kws:
            continue
        if "키워드" in t or t.endswith(":") or t.endswith("："):
            continue  # 머리말/라벨 줄 제외
        kws.append(t)
        if len(kws) >= 6:
            break
    return kws
