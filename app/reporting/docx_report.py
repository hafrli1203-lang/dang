"""당근마켓 광고 문서 생성기 v2.2

광고주 납품 품질을 최우선으로 설계.
templates/sample_report.docx · templates/sample_plan.docx 레이아웃(v1.0)을 재현.

보고서 품질 3대 요소 (document_spec.md v1.0 + v2.2 개선):
  (1) 한 페이지 요약 카드  — _build_summary_card()
        mode별 핵심 KPI 4개를 2행 × 4열 그리드로 시각화
  (2) KPI 표 (모드별)       — _build_kpi_table()
        tracking_mode ∈ {"db_funnel","landing","reaction"} 완전 지원
        _kpi_rows_for_mode() / _kpi_cards_for_mode() 분리
  (3) 차트 3개              — make_charts() → _make_chart_*()
        DPI 200, 흰 배경, 바 값 레이블, CPA 평균선 추가

기타 설계 원칙:
  - 표지: Title 스타일 + Normal 부제/날짜 (templates 스타일 상속)
  - 섹션 헤딩: Heading 1 (#365F91) / Heading 2 (#4F81BD)
  - KPI 표: 2열 (항목명 FFF3E0·Bold | 값), Table Grid
  - 기간별 표: 헤더 FF6F00, 데이터 행 배경 없음
  - 푸터: 표지 제외 전 페이지 — 문서명 + 페이지 번호

공개 API:
    make_charts(timeseries, output_dir, mode)                    → List[Path]
    build_report_docx(meta, kpi, ts, insights, path, chart_dir) → Path
    build_planning_docx(meta, ai_content, path)                 → Path
"""
from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import List, Optional, Tuple

import matplotlib
if not matplotlib.is_interactive():
    matplotlib.use("Agg", force=False)

import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import matplotlib.font_manager as fm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

try:
    from typing import TypedDict
except ImportError:
    from typing_extensions import TypedDict  # type: ignore

logger = logging.getLogger(__name__)

# ── TypedDicts ────────────────────────────────────────────────────────────────

class ProjectMeta(TypedDict, total=False):
    name: str
    period: str
    goal: str
    industry: str
    region: str
    budget: str
    benefits: str           # 주요 혜택 (기획서 광고주 정보 표)
    link: str               # 참고 링크 (기획서 광고주 정보 표)
    campaign_name: str      # 캠페인명 (표지)
    author: str             # 작성자 (표지)
    target: str             # 지역/타겟 상세 (요약카드 메타박스)
    operation_method: str   # 운영방식 (캠페인 개요)


class KPI(TypedDict, total=False):
    total_spend: float
    total_impressions: int
    total_clicks: int
    total_chats: int
    total_followers: int
    total_coupons: int
    ctr: float
    cpc: float
    cpa: float
    # landing mode
    total_conversions: int
    total_link_clicks: int  # 링크클릭 (landing)
    cvr: float
    # reaction mode
    total_reactions: int
    total_likes: int
    total_comments: int
    total_shares: int
    engagement_rate: float
    cpe: float
    # derived
    cost_per_follower: float
    cost_per_coupon: float
    cost_per_chat: float
    status: str             # "좋음"|"주의"|"점검필요"


class TimeseriesRow(TypedDict, total=False):
    date: str
    spend: float
    clicks: int
    chats: int
    impressions: int
    followers: int
    coupons: int
    link_clicks: int    # 링크클릭 (landing mode)
    # landing mode
    conversions: int
    # reaction mode
    reactions: int
    likes: int
    comments: int
    shares: int


class Insights(TypedDict, total=False):
    summary: str
    insights: list
    actions: list


class Experiment(TypedDict, total=False):
    priority: str
    change: str
    success_criteria: str
    owner: str
    schedule: str


class JudgmentCriteria(TypedDict, total=False):
    expand: str
    review: str
    stop: str


class ReportInsights(TypedDict, total=False):
    conclusion: str             # Sec2(C): 결론 3~5줄
    next_actions: list          # Sec2(D): Next Actions
    good: str                   # Sec6: 잘 된 것
    blocked: str                # Sec6: 막힌 것
    hypothesis: str             # Sec6: 가설
    experiments: list           # Sec7: Experiment 리스트
    judgment: JudgmentCriteria  # Sec8: 판단 기준
    # 레거시 호환
    summary: str
    insights: list
    actions: list


# ── Template paths ────────────────────────────────────────────────────────────
# Support both normal and PyInstaller-frozen environments.

def _get_templates_dir() -> Path:
    """Resolve templates directory for both normal and frozen environments."""
    import sys
    if getattr(sys, "frozen", False):
        return Path(sys._MEIPASS) / "templates"
    return Path(__file__).parent.parent.parent / "templates"

_TEMPLATES_DIR   = _get_templates_dir()
_TEMPLATE_REPORT = _TEMPLATES_DIR / "sample_report.docx"
_TEMPLATE_PLAN   = _TEMPLATES_DIR / "sample_plan.docx"

# ── Colour constants ──────────────────────────────────────────────────────────

_C_ORANGE   = "FF6F00"   # 시계열 표 헤더 / 차트 기본색
_C_ORANGE_L = "FFF3E0"   # KPI·정보 표 항목명 열
_C_ORANGE_S = "FFF8F0"   # 요약 카드 값 행 배경 (softest)

_HEX_ORANGE = "#FF6F00"
_HEX_GREEN  = "#43A047"
_HEX_BLUE   = "#1E88E5"
_HEX_PURPLE = "#8E24AA"
_HEX_TEAL   = "#00897B"
_HEX_GRAY   = "#78909C"

_STATUS_COLORS = {
    "좋음":     "4CAF50",   # green
    "주의":     "FF9800",   # orange
    "점검필요": "F44336",   # red
}


# ── Safe value helpers ────────────────────────────────────────────────────────

def _s(val, default: str = "-") -> str:
    if val is None or val == "":
        return default
    return str(val)

def _i(val, default: int = 0) -> int:
    if val is None:
        return default
    try:
        return int(val)
    except (ValueError, TypeError):
        return default

def _f(val, default: float = 0.0) -> float:
    if val is None:
        return default
    try:
        return float(val)
    except (ValueError, TypeError):
        return default


# ── Korean font ───────────────────────────────────────────────────────────────

def _best_korean_font() -> str:
    candidates = ["Malgun Gothic", "맑은 고딕", "NanumGothic", "AppleGothic", "DejaVu Sans"]
    available = {f.name for f in fm.fontManager.ttflist}
    for c in candidates:
        if c in available:
            return c
    return "sans-serif"

_KR_FONT = _best_korean_font()
plt.rcParams["font.family"] = _KR_FONT
plt.rcParams["axes.unicode_minus"] = False


# ── DOCX primitives ───────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _set_cell_padding(cell, pt: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    twips = str(int(pt * 20))
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), twips)
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_width(cell, inches: float) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcW"))
    if existing is not None:
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _remove_table_borders(table) -> None:
    tbl_elm = table._tbl
    tblPr = tbl_elm.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elm.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_table_full_width(table) -> None:
    tbl_elm = table._tbl
    tblPr = tbl_elm.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elm.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblW"))
    if existing is not None:
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)


def _add_page_number_field(run) -> None:
    for tag, text in [
        ("w:fldChar", None),
        ("w:instrText", " PAGE "),
        ("w:fldChar", None),
        ("w:fldChar", None),
    ]:
        el = OxmlElement(tag)
        if tag == "w:fldChar":
            types = ["begin", "separate", "end"]
            idx = [c.tag.split("}")[-1] for c in run._r].count("fldChar")
            el.set(qn("w:fldCharType"), types[min(idx, 2)])
        if text:
            el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)


def _add_footer(doc: Document, label: str) -> None:
    """Footer on all pages except the cover (first page)."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run(f"{label}  ·  ")
    run1.font.size = Pt(8)
    run1.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    run2 = p.add_run()
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
    _add_page_number_field(run2)


def _add_divider(doc: Document, color: str = _C_ORANGE) -> None:
    """Thin horizontal rule (used in _render_md_body for --- lines)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _apply_document_defaults(doc: Document) -> None:
    """A4 page + standard margins (fallback when template file is missing)."""
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)


def _load_template(template_path: Path) -> Document:
    """Load template DOCX (inherits Word styles), clear body content, return fresh doc."""
    doc = Document(str(template_path))
    body = doc.element.body
    sectPr = body.find(qn("w:sectPr"))
    for child in list(body):
        body.remove(child)
    if sectPr is not None:
        body.append(sectPr)
    return doc


# ── Insights normalizer ───────────────────────────────────────────────────────

def _normalize_insights(raw) -> ReportInsights:
    """Convert legacy Insights or new ReportInsights → ReportInsights."""
    if not raw:
        return ReportInsights()
    if raw.get("conclusion"):
        return raw  # already new format
    # legacy → new mapping
    return ReportInsights(
        conclusion=raw.get("summary", ""),
        next_actions=raw.get("actions", []),
        good="",
        blocked="",
        hypothesis="",
        experiments=[],
        judgment={},
        summary=raw.get("summary", ""),
        insights=raw.get("insights", []),
        actions=raw.get("actions", []),
    )


# ── KPI mode mapping ──────────────────────────────────────────────────────────

def _kpi_cards_for_mode(kpi: KPI, mode: str) -> List[Tuple[str, str]]:
    """한 페이지 요약 카드용 핵심 KPI 4개 반환 (label, value).

    tracking_mode:
        "db_funnel" → 광고비 / CTR / CPC / 총문의
        "landing"   → 광고비 / CTR / CVR / CPA
        "reaction"  → 광고비 / 총노출 / 총반응 / CPE
    """
    if mode == "landing":
        return [
            ("총 광고비용",    f"{_i(kpi.get('total_spend')):,} 원"),
            ("클릭률  CTR",    f"{_f(kpi.get('ctr')):.2f} %"),
            ("전환율  CVR",    f"{_f(kpi.get('cvr')):.2f} %"),
            ("전환당비용 CPA", f"{_f(kpi.get('cpa')):,.0f} 원"),
        ]
    if mode == "reaction":
        return [
            ("총 광고비용",    f"{_i(kpi.get('total_spend')):,} 원"),
            ("총 노출수",      f"{_i(kpi.get('total_impressions')):,} 회"),
            ("총 반응수",      f"{_i(kpi.get('total_reactions')):,} 건"),
            ("반응당비용 CPE", f"{_f(kpi.get('cpe')):,.0f} 원"),
        ]
    # db_funnel (default)
    return [
        ("총 광고비용",    f"{_i(kpi.get('total_spend')):,} 원"),
        ("클릭률  CTR",    f"{_f(kpi.get('ctr')):.2f} %"),
        ("클릭당비용 CPC", f"{_f(kpi.get('cpc')):,.0f} 원"),
        ("총 문의",        f"{_i(kpi.get('total_chats')):,} 건"),
    ]


def _kpi_rows_for_mode(kpi: KPI, mode: str) -> List[Tuple[str, str, str]]:
    """KPI 상세 표용 전체 행 반환 (metric, value, note).

    tracking_mode:
        "db_funnel" → 9행 (기본 소상공인 퍼널)
        "landing"   → 8행 (랜딩페이지 전환)
        "reaction"  → 8행 (콘텐츠 반응)
    """
    if mode == "landing":
        return [
            ("총 광고 비용",    f"{_i(kpi.get('total_spend')):,} 원",        "집행 예산"),
            ("총 노출 수",      f"{_i(kpi.get('total_impressions')):,} 회",  "도달 범위"),
            ("총 클릭 수",      f"{_i(kpi.get('total_clicks')):,} 회",       ""),
            ("클릭률 (CTR)",    f"{_f(kpi.get('ctr')):.2f} %",               ""),
            ("클릭당비용 (CPC)", f"{_f(kpi.get('cpc')):,.0f} 원",           "낮을수록 효율적"),
            ("총 전환 수",      f"{_i(kpi.get('total_conversions')):,} 건",  "랜딩 전환"),
            ("전환율 (CVR)",    f"{_f(kpi.get('cvr')):.2f} %",               "클릭→전환"),
            ("전환당비용 (CPA)", f"{_f(kpi.get('cpa')):,.0f} 원",           "낮을수록 효율적"),
        ]
    if mode == "reaction":
        er_val = kpi.get("engagement_rate") or _f(kpi.get("ctr", 0.0))
        return [
            ("총 광고 비용",   f"{_i(kpi.get('total_spend')):,} 원",        ""),
            ("총 노출 수",     f"{_i(kpi.get('total_impressions')):,} 회",  "도달 범위"),
            ("총 반응 수",     f"{_i(kpi.get('total_reactions')):,} 건",    "좋아요+댓글+공유"),
            ("좋아요",         f"{_i(kpi.get('total_likes')):,} 건",        ""),
            ("댓글",           f"{_i(kpi.get('total_comments')):,} 건",     ""),
            ("공유",           f"{_i(kpi.get('total_shares')):,} 건",       ""),
            ("반응률 (ER)",    f"{_f(er_val):.2f} %",                       "반응÷노출"),
            ("반응당비용 (CPE)", f"{_f(kpi.get('cpe')):,.0f} 원",          "낮을수록 효율적"),
        ]
    # db_funnel (default)
    return [
        ("총 광고 비용",    f"{_i(kpi.get('total_spend')):,} 원",        "집행 예산"),
        ("총 노출 수",      f"{_i(kpi.get('total_impressions')):,} 회",  "도달 범위"),
        ("총 클릭 수",      f"{_i(kpi.get('total_clicks')):,} 회",       ""),
        ("클릭률 (CTR)",    f"{_f(kpi.get('ctr')):.2f} %",               "업종 평균 약 2.5%"),
        ("클릭당 비용 (CPC)", f"{_f(kpi.get('cpc')):,.0f} 원",          "낮을수록 효율적"),
        ("총 문의 수",      f"{_i(kpi.get('total_chats')):,} 건",        "채팅 문의"),
        ("문의당 비용 (CPA)", f"{_f(kpi.get('cpa')):,.0f} 원",          "낮을수록 효율적"),
        ("단골 전환",       f"{_i(kpi.get('total_followers')):,} 명",    ""),
        ("쿠폰 사용",       f"{_i(kpi.get('total_coupons')):,} 건",      ""),
    ]


# ── Chart helpers ─────────────────────────────────────────────────────────────

def _style_axes(ax, title: str, ylabel: str = "") -> None:
    """공통 축 스타일 — 스파인 제거, 그리드, 타이틀."""
    ax.set_title(title, fontsize=13, fontweight="bold", pad=12, color="#37474F")
    if ylabel:
        ax.set_ylabel(ylabel, fontsize=9, color=_HEX_GRAY)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#DDDDDD")
    ax.spines["bottom"].set_color("#DDDDDD")
    ax.yaxis.grid(True, linestyle="--", alpha=0.35, color="#EEEEEE", zorder=0)
    ax.set_axisbelow(True)
    ax.tick_params(colors=_HEX_GRAY, labelsize=9)


def _white_bg(fig, ax) -> None:
    """그림·축 배경을 순백으로 고정 (PDF/Word 삽입 시 투명 배경 방지)."""
    fig.patch.set_facecolor("white")
    ax.set_facecolor("#FAFAFA")


def _annotate_bars(ax, x_vals, y_vals, fmt="{:,.0f}", color: str = "#37474F",
                   offset_pct: float = 0.03) -> None:
    """막대 위에 값 레이블 추가. fmt은 str ('{:,.0f}') 또는 callable (lambda v: ...) 가능."""
    max_v = max((abs(v) for v in y_vals), default=1) or 1
    for xi, yi in zip(x_vals, y_vals):
        if yi:
            label = fmt(yi) if callable(fmt) else fmt.format(yi)
            ax.text(xi, yi + max_v * offset_pct, label,
                    ha="center", va="bottom", fontsize=8, color=color, fontweight="bold")


def _save_chart(fig, path: Path, dpi: int = 200) -> Path:
    try:
        fig.savefig(path, dpi=dpi, bbox_inches="tight", facecolor="white")
    except MemoryError:
        # bbox_inches="tight" can cause very large bitmaps; retry without it
        fig.savefig(path, dpi=min(dpi, 150), facecolor="white")
    plt.close("all")
    return path


# ── Mode-aware chart builders ─────────────────────────────────────────────────

def _make_chart_spend_action(
    labels: List[str],
    spends: List[float],
    timeseries: List[TimeseriesRow],
    mode: str,
    output_dir: Path,
) -> Optional[Path]:
    """Chart 1: 기간별 광고비용(막대) vs 주요 행동 지표(꺾은선) — dual-axis.

    mode="db_funnel": 행동 = 클릭·문의
    mode="landing":   행동 = 클릭·전환
    mode="reaction":  행동 = 반응수
    """
    try:
        fig, ax1 = plt.subplots(figsize=(8, 4))
        _white_bg(fig, ax1)
        x = list(range(len(labels)))

        bars = ax1.bar(x, spends, color=_HEX_ORANGE, alpha=0.80,
                       label="광고비용", zorder=2, width=0.55)
        ax1.set_ylabel("비용 (원)", fontsize=9, color=_HEX_ORANGE)
        ax1.tick_params(axis="y", labelcolor=_HEX_ORANGE, labelsize=9)
        ax1.spines["top"].set_visible(False)
        ax1.yaxis.grid(True, linestyle="--", alpha=0.25, zorder=0)
        ax1.set_axisbelow(True)
        ax1.yaxis.set_major_formatter(mticker.FuncFormatter(
            lambda v, _: f"{int(v)//1000}K" if v >= 1000 else f"{int(v)}"
        ))

        # 막대 위 값 레이블 (단위: 만원)
        _annotate_bars(ax1, x, spends,
                       fmt=lambda v: f"{v/10000:.1f}만" if v >= 10000 else f"{v:,.0f}",
                       color=_HEX_ORANGE, offset_pct=0.025)

        ax2 = ax1.twinx()
        ax2.spines["top"].set_visible(False)
        ax2.spines["right"].set_color("#DDDDDD")
        ax2.tick_params(axis="y", labelsize=9, labelcolor=_HEX_BLUE)

        if mode == "landing":
            clicks = [_i(r.get("clicks")) for r in timeseries]
            convs  = [_i(r.get("conversions")) for r in timeseries]
            ax2.plot(x, clicks, "o-", color=_HEX_GREEN,  label="클릭", lw=2, markersize=6, zorder=3)
            ax2.plot(x, convs,  "s-", color=_HEX_BLUE,   label="전환", lw=2, markersize=6, zorder=3)
            ax2.set_ylabel("클릭 / 전환 (건)", fontsize=9, color=_HEX_BLUE)
        elif mode == "reaction":
            reactions = [
                _i(r.get("reactions")) or
                _i(r.get("likes")) + _i(r.get("comments")) + _i(r.get("shares"))
                for r in timeseries
            ]
            ax2.plot(x, reactions, "o-", color=_HEX_PURPLE, label="반응", lw=2, markersize=6, zorder=3)
            ax2.set_ylabel("반응 수 (건)", fontsize=9, color=_HEX_PURPLE)
            ax2.tick_params(axis="y", labelcolor=_HEX_PURPLE)
        else:  # db_funnel
            clicks = [_i(r.get("clicks")) for r in timeseries]
            chats  = [_i(r.get("chats"))  for r in timeseries]
            ax2.plot(x, clicks, "o-", color=_HEX_GREEN, label="클릭", lw=2, markersize=6, zorder=3)
            ax2.plot(x, chats,  "s-", color=_HEX_BLUE,  label="문의", lw=2, markersize=6, zorder=3)
            ax2.set_ylabel("클릭 / 문의 (건)", fontsize=9, color=_HEX_BLUE)

        ax1.set_xticks(x)
        ax1.set_xticklabels(labels, rotation=20, ha="right", fontsize=9)
        ax1.set_title("기간별 광고비용 vs 행동 지표", fontsize=13,
                      fontweight="bold", pad=12, color="#37474F")

        h1, l1 = ax1.get_legend_handles_labels()
        h2, l2 = ax2.get_legend_handles_labels()
        ax1.legend(h1 + h2, l1 + l2, fontsize=8, framealpha=0.7,
                   loc="upper left", frameon=True)

        plt.tight_layout()
        return _save_chart(fig, output_dir / "chart_spend_action.png")
    except Exception:
        logger.exception("chart_spend_action failed")
        plt.close("all")
        return None


def _make_chart_cpa_trend(
    labels: List[str],
    spends: List[float],
    timeseries: List[TimeseriesRow],
    mode: str,
    output_dir: Path,
) -> Optional[Path]:
    """Chart 2: 기간별 CPA / CPE 비용 효율 추이 (꺾은선 + 면 + 평균선).

    mode="db_funnel": 문의당 비용 (CPA)
    mode="landing":   전환당 비용 (CPA)
    mode="reaction":  반응당 비용 (CPE)
    """
    try:
        if mode == "landing":
            divisors = [_i(r.get("conversions")) for r in timeseries]
            ylabel, title, line_label = "CPA (원)", "기간별 전환당 비용 (CPA)", "CPA"
            line_color = _HEX_BLUE
        elif mode == "reaction":
            divisors = [
                _i(r.get("reactions")) or
                _i(r.get("likes")) + _i(r.get("comments")) + _i(r.get("shares"))
                for r in timeseries
            ]
            ylabel, title, line_label = "CPE (원)", "기간별 반응당 비용 (CPE)", "CPE"
            line_color = _HEX_PURPLE
        else:  # db_funnel
            divisors = [_i(r.get("chats")) for r in timeseries]
            ylabel, title, line_label = "CPA (원)", "기간별 문의당 비용 (CPA)", "CPA"
            line_color = _HEX_PURPLE

        values = [s / d if d > 0 else 0.0 for s, d in zip(spends, divisors)]

        fig, ax = plt.subplots(figsize=(8, 4))
        _white_bg(fig, ax)

        ax.plot(labels, values, marker="o", color=line_color, linewidth=2.5,
                markersize=8, markerfacecolor="white", markeredgewidth=2.5,
                label=line_label, zorder=4)
        ax.fill_between(labels, values, alpha=0.08, color=line_color)

        # 각 포인트 값 레이블
        max_v = max(values, default=0.1) or 0.1
        for i, v in enumerate(values):
            if v > 0:
                ax.text(i, v + max_v * 0.05, f"{v:,.0f}",
                        ha="center", va="bottom", fontsize=8,
                        color=line_color, fontweight="bold")

        # 평균선 (유효값만)
        valid = [v for v in values if v > 0]
        if valid:
            avg = sum(valid) / len(valid)
            ax.axhline(avg, color=_HEX_GRAY, linestyle="--", linewidth=1.2,
                       alpha=0.8, label=f"평균 {avg:,.0f}원", zorder=2)

        _style_axes(ax, title, ylabel)
        ax.tick_params(axis="x", rotation=20)
        ax.legend(fontsize=8, framealpha=0.7)
        plt.tight_layout()
        return _save_chart(fig, output_dir / "chart_cpa.png")
    except Exception:
        logger.exception("chart_cpa_trend failed")
        plt.close("all")
        return None


def _make_chart_funnel(
    timeseries: List[TimeseriesRow],
    mode: str,
    output_dir: Path,
) -> Optional[Path]:
    """Chart 3: 전환 퍼널 수평 막대 (기간 합산).

    mode="db_funnel": 노출→클릭→문의→단골
    mode="landing":   노출→클릭→전환
    mode="reaction":  노출→반응→좋아요·댓글·공유
    """
    try:
        if mode == "landing":
            stages = [
                ("노출", sum(_i(r.get("impressions")) for r in timeseries)),
                ("클릭", sum(_i(r.get("clicks"))      for r in timeseries)),
                ("전환", sum(_i(r.get("conversions")) for r in timeseries)),
            ]
            colors = [_HEX_ORANGE, _HEX_GREEN, _HEX_BLUE]
            title  = "전환 퍼널 (노출 → 클릭 → 전환)"
        elif mode == "reaction":
            total_likes    = sum(_i(r.get("likes"))    for r in timeseries)
            total_comments = sum(_i(r.get("comments")) for r in timeseries)
            total_shares   = sum(_i(r.get("shares"))   for r in timeseries)
            total_reac     = sum(
                _i(r.get("reactions")) or
                _i(r.get("likes")) + _i(r.get("comments")) + _i(r.get("shares"))
                for r in timeseries
            )
            stages = [
                ("노출",   sum(_i(r.get("impressions")) for r in timeseries)),
                ("반응",   total_reac),
                ("좋아요", total_likes),
                ("댓글",   total_comments),
                ("공유",   total_shares),
            ]
            colors = [_HEX_ORANGE, _HEX_PURPLE, "#E91E63", "#9C27B0", "#3F51B5"]
            title  = "반응 퍼널 (노출 → 반응 분해)"
        else:  # db_funnel
            stages = [
                ("노출",     sum(_i(r.get("impressions")) for r in timeseries)),
                ("클릭",     sum(_i(r.get("clicks"))      for r in timeseries)),
                ("문의",     sum(_i(r.get("chats"))       for r in timeseries)),
                ("단골전환", sum(_i(r.get("followers"))   for r in timeseries)),
            ]
            colors = [_HEX_ORANGE, _HEX_GREEN, _HEX_BLUE, _HEX_TEAL]
            title  = "전환 퍼널 (노출 → 클릭 → 문의 → 단골)"

        max_val = stages[0][1] if stages else 1
        if max_val == 0:
            max_val = 1
        stage_labels = [s[0] for s in stages]
        values_f     = [s[1] for s in stages]
        n            = len(stages)
        y_pos        = list(range(n - 1, -1, -1))  # 위→아래

        fig, ax = plt.subplots(figsize=(8, min(6, max(3.5, n * 0.9))))
        _white_bg(fig, ax)

        for yp, val, color in zip(y_pos, values_f, colors):
            ax.barh(yp, val, color=color, alpha=0.85, height=0.55, zorder=3)
            pct = val / max_val * 100
            label_x = val + max_val * 0.012
            ax.text(label_x, yp,
                    f"{val:,}  ({pct:.1f}%)",
                    va="center", ha="left", fontsize=9, color="#37474F")

        ax.set_yticks(y_pos)
        ax.set_yticklabels(stage_labels, fontsize=10, fontweight="bold")
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_color("#DDDDDD")
        ax.spines["bottom"].set_color("#DDDDDD")
        ax.xaxis.grid(True, linestyle="--", alpha=0.3, zorder=0)
        ax.set_axisbelow(True)
        ax.tick_params(colors=_HEX_GRAY, labelsize=9)
        ax.xaxis.set_major_formatter(mticker.FuncFormatter(
            lambda v, _: f"{int(v)//1000}K" if v >= 1000 else f"{int(v)}"
        ))
        ax.set_title(title, fontsize=13, fontweight="bold", pad=12, color="#37474F")
        # x축 오른쪽 여백 확보 (레이블 잘림 방지)
        ax.set_xlim(right=max_val * 1.35)
        plt.tight_layout()
        return _save_chart(fig, output_dir / "chart_funnel.png")
    except Exception:
        logger.exception("chart_funnel failed")
        plt.close("all")
        return None


# ── Section builders ──────────────────────────────────────────────────────────

def _build_cover(
    doc: Document,
    project_meta: ProjectMeta,
    doc_title: str,
    is_planning: bool = False,
) -> None:
    """표지: Title 스타일 + Normal 부제/날짜 + 캠페인명/작성자 + 신뢰 안내문."""
    try:
        doc.add_paragraph(doc_title, style="Title")

        name     = _s(project_meta.get("name"), "")
        industry = _s(project_meta.get("industry"), "")
        region   = _s(project_meta.get("region"), "")
        if is_planning:
            parts = [x for x in [name, industry, region] if x and x != "-"]
        else:
            parts = [x for x in [name, region] if x and x != "-"]
        subtitle = "  |  ".join(parts)
        if subtitle:
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_sub.add_run(subtitle)
            run.font.size = Pt(13)
            if is_planning:
                run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

        # 캠페인명 / 작성자 (성과보고서 전용)
        if not is_planning:
            campaign = _s(project_meta.get("campaign_name"), "")
            author = _s(project_meta.get("author"), "")
            if campaign and campaign != "-":
                p_camp = doc.add_paragraph()
                p_camp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_camp.add_run(f"캠페인: {campaign}")
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
            if author and author != "-":
                p_auth = doc.add_paragraph()
                p_auth.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_auth.add_run(f"작성자: {author}")
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

        date_label = "작성일" if is_planning else "보고서 작성일"
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_date.add_run(f"{date_label}: {datetime.now().strftime('%Y년 %m월 %d일')}")

        doc.add_paragraph()

        # 신뢰 안내문 (성과보고서 전용)
        if not is_planning:
            for notice in [
                "수치 해석은 추적 환경에 따라 달라질 수 있습니다.",
                "과장/단정 표현 없이, 확인 가능한 데이터 기준으로만 제안합니다.",
            ]:
                p_n = doc.add_paragraph()
                p_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_n.add_run(notice)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    except Exception:
        logger.exception("_build_cover failed")
        doc.add_paragraph("[표지 생성 오류]")


def _build_summary_card(doc: Document, kpi: KPI, mode: str = "db_funnel") -> None:
    """(1) 한 페이지 요약 카드 — mode별 핵심 KPI 4개, 2행 × 4열 그리드.

    상단 행: 지표명 (FF6F00 배경, 흰색 Bold 8.5pt)
    하단 행: 값     (FFF3E0 배경, 주황 Bold 17pt)
    tracking_mode별 카드 조합은 _kpi_cards_for_mode() 참조.
    """
    try:
        cards = _kpi_cards_for_mode(kpi, mode)
        tbl = doc.add_table(rows=2, cols=4)
        tbl.style = "Table Grid"
        _set_table_full_width(tbl)

        for i, (label, value) in enumerate(cards):
            # ── 상단: 지표명 ──────────────────────────────────────
            lc = tbl.rows[0].cells[i]
            lc.paragraphs[0].clear()
            lc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = lc.paragraphs[0].add_run(label)
            run.font.size = Pt(8.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            _set_cell_bg(lc, _C_ORANGE)
            _set_cell_padding(lc, 6)

            # ── 하단: 값 ─────────────────────────────────────────
            vc = tbl.rows[1].cells[i]
            vc.paragraphs[0].clear()
            vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = vc.paragraphs[0].add_run(value)
            run.font.size = Pt(17)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x6F, 0x00)
            _set_cell_bg(vc, _C_ORANGE_L)
            _set_cell_padding(vc, 10)

    except Exception:
        logger.exception("_build_summary_card failed")
        doc.add_paragraph("[요약 카드 생성 오류]")


def _build_kpi_table(doc: Document, kpi: KPI, mode: str = "db_funnel") -> None:
    """(2) KPI 상세 표 — 2열 (항목명 FFF3E0·Bold | 값), Table Grid.

    tracking_mode별 행 구성은 _kpi_rows_for_mode() 참조.
    항목명 열 고정폭 2.1 in, 값 열 나머지.
    """
    try:
        rows_data = _kpi_rows_for_mode(kpi, mode)
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        _set_table_full_width(tbl)

        for metric, value, _ in rows_data:
            row = tbl.add_row()

            # 항목명 열 (고정폭 + FFF3E0 배경 + Bold)
            lc = row.cells[0]
            lc.text = metric
            _set_cell_bg(lc, _C_ORANGE_L)
            _set_cell_padding(lc, 5)
            _set_cell_width(lc, 2.1)
            for r in lc.paragraphs[0].runs:
                r.bold = True
                r.font.size = Pt(10)

            # 값 열
            vc = row.cells[1]
            vc.text = value
            _set_cell_padding(vc, 5)
            for r in vc.paragraphs[0].runs:
                r.font.size = Pt(10)

    except Exception:
        logger.exception("_build_kpi_table failed")
        doc.add_paragraph("[KPI 표 생성 오류]")


def _build_timeseries_table(doc: Document, timeseries: List[TimeseriesRow]) -> None:
    """기간별 성과 표 — 헤더 FF6F00, 데이터 행 배경 없음, 7열."""
    try:
        if not timeseries:
            doc.add_paragraph("기간별 데이터가 없습니다.")
            return
        cols = ["기간", "비용(원)", "노출", "클릭", "문의", "단골", "쿠폰"]
        tbl = doc.add_table(rows=1, cols=len(cols))
        tbl.style = "Table Grid"
        _set_table_full_width(tbl)

        for i, col in enumerate(cols):
            cell = tbl.rows[0].cells[i]
            cell.text = col
            _set_cell_bg(cell, _C_ORANGE)
            _set_cell_padding(cell, 5)
            for r in cell.paragraphs[0].runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)

        for row in timeseries:
            dr = tbl.add_row()
            vals = [
                _s(row.get("date"), ""),
                f"{_i(row.get('spend')):,}",
                f"{_i(row.get('impressions')):,}",
                f"{_i(row.get('clicks')):,}",
                f"{_i(row.get('chats')):,}",
                f"{_i(row.get('followers')):,}",
                f"{_i(row.get('coupons')):,}",
            ]
            for i, v in enumerate(vals):
                cell = dr.cells[i]
                cell.text = v
                _set_cell_padding(cell, 4)
                for r in cell.paragraphs[0].runs:
                    r.font.size = Pt(9)

    except Exception:
        logger.exception("_build_timeseries_table failed")
        doc.add_paragraph("[기간별 성과 표 생성 오류]")


def _build_charts_section(doc: Document, chart_paths: List[Path]) -> None:
    """(3) 차트 3개 삽입 — 가운데 정렬 + 이탤릭 캡션."""
    try:
        if not chart_paths:
            doc.add_paragraph("생성된 차트가 없습니다.")
            return
        captions = {
            "chart_spend_action.png": "Fig 1.  기간별 광고비용 vs 행동 지표",
            "chart_cpa.png":          "Fig 2.  기간별 비용 효율 (CPA / CPE) 추이",
            "chart_funnel.png":       "Fig 3.  전환 퍼널 분석",
        }
        for cp in chart_paths:
            if not cp.exists():
                continue
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(str(cp), width=Inches(5.5))

            caption = captions.get(cp.name, cp.stem)
            p_cap = doc.add_paragraph(caption)
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p_cap.runs:
                r.font.size = Pt(8)
                r.font.italic = True
                r.font.color.rgb = RGBColor(0x78, 0x78, 0x78)
            doc.add_paragraph()

    except Exception:
        logger.exception("_build_charts_section failed")
        doc.add_paragraph("[차트 섹션 생성 오류]")


def _build_toc(doc: Document) -> None:
    """목차 (Sec 1) — Word TOC 필드 코드 삽입."""
    try:
        doc.add_heading("목차", level=1)
        p = doc.add_paragraph()
        run = p.add_run()
        r = run._r

        # BEGIN field
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        r.append(fld_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-2" \\h \\z \\u '
        r.append(instr)

        # SEPARATE field
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        r.append(fld_sep)

        # END field
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        r.append(fld_end)

        p_hint = doc.add_paragraph()
        p_hint.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_hint.add_run("(Word에서 이 문서를 열고 Ctrl+A → F9를 누르면 목차가 갱신됩니다)")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True

    except Exception:
        logger.exception("_build_toc failed")
        doc.add_paragraph("[목차 생성 오류]")


def _build_meta_box(doc: Document, project_meta: ProjectMeta, kpi: KPI) -> None:
    """메타 박스 (Sec 2B) — 2×2 borderless 테이블."""
    try:
        tbl = doc.add_table(rows=2, cols=2)
        _remove_table_borders(tbl)
        _set_table_full_width(tbl)

        goal = _s(project_meta.get("goal"), "-")
        status = _s(kpi.get("status"), "")
        status_label = f"판단 KPI: {status}" if status else "판단 KPI: -"
        target = _s(project_meta.get("target"), _s(project_meta.get("region"), "-"))
        benefits = _s(project_meta.get("benefits"), "-")

        items = [
            ("목표", goal),
            (status_label.split(":")[0], status_label.split(":")[-1].strip()),
            ("타겟·지역", target),
            ("주요 혜택", benefits),
        ]

        for idx, (label, value) in enumerate(items):
            row_i = idx // 2
            col_i = idx % 2
            cell = tbl.rows[row_i].cells[col_i]
            cell.paragraphs[0].clear()
            run_l = cell.paragraphs[0].add_run(f"{label}: ")
            run_l.font.bold = True
            run_l.font.size = Pt(9)
            run_l.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
            run_v = cell.paragraphs[0].add_run(value)
            run_v.font.size = Pt(9)
            _set_cell_padding(cell, 4)

    except Exception:
        logger.exception("_build_meta_box failed")
        doc.add_paragraph("[메타 박스 생성 오류]")


def _build_conclusion(doc: Document, insights: ReportInsights) -> None:
    """결론 (Sec 2C) — insights.conclusion을 Heading 2 + Normal 단락들로 출력."""
    try:
        conclusion = _s(insights.get("conclusion"), "")
        if not conclusion or conclusion == "-":
            return
        doc.add_heading("결론", level=2)
        for line in conclusion.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
    except Exception:
        logger.exception("_build_conclusion failed")
        doc.add_paragraph("[결론 생성 오류]")


def _build_next_actions(doc: Document, insights: ReportInsights) -> None:
    """Next Actions (Sec 2D) — checkbox 형태, 최대 7개."""
    try:
        actions = insights.get("next_actions") or []
        if not actions:
            return
        doc.add_heading("Next Actions", level=2)
        for item in actions[:7]:
            text = _s(item, "").strip()
            if text and text != "-":
                doc.add_paragraph(f"\u2610  {text}")
    except Exception:
        logger.exception("_build_next_actions failed")
        doc.add_paragraph("[Next Actions 생성 오류]")


def _build_campaign_overview(
    doc: Document, project_meta: ProjectMeta, tracking_mode: str = "db_funnel",
) -> None:
    """캠페인 개요 (Sec 3) — 5행×2열 Table Grid."""
    try:
        mode_labels = {
            "db_funnel": "DB 퍼널 (노출→클릭→문의→단골)",
            "landing": "랜딩 페이지 전환",
            "reaction": "콘텐츠 반응 (좋아요·댓글·공유)",
        }
        fields = [
            ("목적", _s(project_meta.get("goal"), "-")),
            ("기간", _s(project_meta.get("period"), "-")),
            ("예산", _s(project_meta.get("budget"), "-")),
            ("운영방식", _s(project_meta.get("operation_method"), "-")),
            ("추적모드", mode_labels.get(tracking_mode, tracking_mode)),
        ]
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        _set_table_full_width(tbl)

        for label, value in fields:
            row = tbl.add_row()
            lc = row.cells[0]
            lc.text = label
            _set_cell_bg(lc, _C_ORANGE_L)
            _set_cell_padding(lc, 5)
            _set_cell_width(lc, 2.0)
            for r in lc.paragraphs[0].runs:
                r.bold = True
                r.font.size = Pt(10)
            vc = row.cells[1]
            vc.text = value
            _set_cell_padding(vc, 5)
            for r in vc.paragraphs[0].runs:
                r.font.size = Pt(10)

    except Exception:
        logger.exception("_build_campaign_overview failed")
        doc.add_paragraph("[캠페인 개요 생성 오류]")


def _build_insights_section(doc: Document, insights: ReportInsights) -> None:
    """인사이트 (Sec 6) — good/blocked/hypothesis 각 1~3줄."""
    try:
        good = _s(insights.get("good"), "")
        blocked = _s(insights.get("blocked"), "")
        hypothesis = _s(insights.get("hypothesis"), "")

        has_new = any(v and v != "-" for v in [good, blocked, hypothesis])

        if has_new:
            if good and good != "-":
                doc.add_heading("잘 된 것", level=2)
                for line in good.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())

            if blocked and blocked != "-":
                doc.add_heading("막힌 것", level=2)
                for line in blocked.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())

            if hypothesis and hypothesis != "-":
                doc.add_heading("가설", level=2)
                for line in hypothesis.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())
        else:
            # fallback: 레거시 insights 리스트
            legacy = insights.get("insights") or []
            if legacy:
                for i, item in enumerate(legacy, 1):
                    p = doc.add_paragraph()
                    run = p.add_run(f"인사이트 {i}  \u2013  {_s(item)}")
                    run.font.bold = True
            else:
                doc.add_paragraph("인사이트 정보 없음")

    except Exception:
        logger.exception("_build_insights_section failed")
        doc.add_paragraph("[인사이트 섹션 생성 오류]")


def _build_experiments_table(doc: Document, insights: ReportInsights) -> None:
    """실험·개선안 (Sec 7) — 5열 Table Grid, 오렌지 헤더."""
    try:
        experiments = insights.get("experiments") or []

        # fallback: 레거시 actions → change 열에 넣기
        if not experiments:
            legacy_actions = insights.get("actions") or []
            if legacy_actions:
                experiments = [
                    {"priority": str(i), "change": _s(a), "success_criteria": "-",
                     "owner": "-", "schedule": "-"}
                    for i, a in enumerate(legacy_actions, 1)
                ]

        if not experiments:
            doc.add_paragraph("실험·개선안 없음")
            return

        headers = ["우선순위", "변경 내용", "성공 기준", "담당", "일정"]
        tbl = doc.add_table(rows=1, cols=5)
        tbl.style = "Table Grid"
        _set_table_full_width(tbl)

        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, _C_ORANGE)
            _set_cell_padding(cell, 5)
            for r in cell.paragraphs[0].runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(9)

        for exp in experiments:
            if isinstance(exp, dict):
                vals = [
                    _s(exp.get("priority"), "-"),
                    _s(exp.get("change"), "-"),
                    _s(exp.get("success_criteria"), "-"),
                    _s(exp.get("owner"), "-"),
                    _s(exp.get("schedule"), "-"),
                ]
            else:
                vals = ["-", _s(exp), "-", "-", "-"]
            row = tbl.add_row()
            for i, v in enumerate(vals):
                cell = row.cells[i]
                cell.text = v
                _set_cell_padding(cell, 4)
                for r in cell.paragraphs[0].runs:
                    r.font.size = Pt(9)

    except Exception:
        logger.exception("_build_experiments_table failed")
        doc.add_paragraph("[실험·개선안 생성 오류]")


def _build_judgment_criteria(doc: Document, insights: ReportInsights) -> None:
    """판단 기준 (Sec 8) — 3행×2열 Table Grid: 확대/검토/중단."""
    try:
        judgment = insights.get("judgment") or {}
        if not judgment:
            doc.add_paragraph("판단 기준 없음")
            return

        labels = [
            ("확대", judgment.get("expand", "-")),
            ("검토", judgment.get("review", "-")),
            ("중단", judgment.get("stop", "-")),
        ]

        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        _set_table_full_width(tbl)

        for label, value in labels:
            row = tbl.add_row()
            lc = row.cells[0]
            lc.text = label
            _set_cell_bg(lc, _C_ORANGE_L)
            _set_cell_padding(lc, 5)
            _set_cell_width(lc, 1.5)
            for r in lc.paragraphs[0].runs:
                r.bold = True
                r.font.size = Pt(10)
            vc = row.cells[1]
            vc.text = _s(value)
            _set_cell_padding(vc, 5)
            for r in vc.paragraphs[0].runs:
                r.font.size = Pt(10)

    except Exception:
        logger.exception("_build_judgment_criteria failed")
        doc.add_paragraph("[판단 기준 생성 오류]")


def _build_appendix(
    doc: Document, timeseries: List[TimeseriesRow], mode: str = "db_funnel",
) -> None:
    """부록 (Sec 9) — 기간별 성과 표 + 지표 정의 표."""
    try:
        # 기간별 성과 표 (최대 10행)
        if timeseries:
            doc.add_heading("기간별 원본 데이터", level=2)
            _build_timeseries_table(doc, timeseries[:10])
            doc.add_paragraph()

        # 지표 정의 표
        doc.add_heading("지표 정의", level=2)

        if mode == "landing":
            defs = [
                ("CTR", "클릭률 = 클릭 ÷ 노출 × 100"),
                ("CPC", "클릭당 비용 = 광고비 ÷ 클릭"),
                ("CVR", "전환율 = 전환 ÷ 클릭 × 100"),
                ("CPA", "전환당 비용 = 광고비 ÷ 전환"),
            ]
        elif mode == "reaction":
            defs = [
                ("노출", "광고가 사용자 화면에 표시된 횟수"),
                ("반응", "좋아요 + 댓글 + 공유의 합계"),
                ("ER", "반응률 = 반응 ÷ 노출 × 100"),
                ("CPE", "반응당 비용 = 광고비 ÷ 반응"),
            ]
        else:  # db_funnel
            defs = [
                ("CTR", "클릭률 = 클릭 ÷ 노출 × 100"),
                ("CPC", "클릭당 비용 = 광고비 ÷ 클릭"),
                ("CPA", "문의당 비용 = 광고비 ÷ 문의"),
                ("단골", "광고를 통해 단골(팔로워)로 전환된 수"),
            ]

        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        _set_table_full_width(tbl)

        for abbr, definition in defs:
            row = tbl.add_row()
            lc = row.cells[0]
            lc.text = abbr
            _set_cell_bg(lc, _C_ORANGE_L)
            _set_cell_padding(lc, 4)
            _set_cell_width(lc, 1.2)
            for r in lc.paragraphs[0].runs:
                r.bold = True
                r.font.size = Pt(9)
            vc = row.cells[1]
            vc.text = definition
            _set_cell_padding(vc, 4)
            for r in vc.paragraphs[0].runs:
                r.font.size = Pt(9)

    except Exception:
        logger.exception("_build_appendix failed")
        doc.add_paragraph("[부록 생성 오류]")


def _build_ai_analysis(doc: Document, insights: Insights) -> None:
    """AI 성과 분석 — Heading 1 + 3개의 Heading 2 서브섹션 (document_spec.md v1.0)."""
    try:
        doc.add_heading("AI 성과 분석", level=1)

        doc.add_heading("1. 성과 요약", level=2)
        summary = _s(insights.get("summary"), "")
        if summary and summary != "-":
            for line in summary.split("\n"):
                if line.strip():
                    doc.add_paragraph(line.strip())

        doc.add_heading("2. 주요 인사이트", level=2)
        insight_list = insights.get("insights") or []
        if insight_list:
            for i, item in enumerate(insight_list, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"인사이트 {i}  \u2013  {_s(item)}")
                run.font.bold = True
        else:
            doc.add_paragraph("인사이트 정보 없음")

        doc.add_heading("3. 다음 액션 추천", level=2)
        action_list = insights.get("actions") or []
        if action_list:
            for i, item in enumerate(action_list, 1):
                doc.add_paragraph(f"{i}.  {_s(item)}")
        else:
            doc.add_paragraph("액션 정보 없음")

    except Exception:
        logger.exception("_build_ai_analysis failed")
        doc.add_paragraph("[AI 분석 섹션 생성 오류]")


def _build_advertiser_profile(doc: Document, project_meta: ProjectMeta) -> None:
    """광고주 정보 섹션 — Heading 1 + 8행 × 2열 표 (document_spec.md v1.0)."""
    try:
        doc.add_heading("광고주 정보", level=1)
        doc.add_paragraph()

        fields = [
            ("광고주명",  project_meta.get("name")),
            ("업종",      project_meta.get("industry")),
            ("지역",      project_meta.get("region")),
            ("광고 목표", project_meta.get("goal")),
            ("예산",      project_meta.get("budget")),
            ("집행 기간", project_meta.get("period")),
            ("주요 혜택", project_meta.get("benefits", "-")),
            ("참고 링크", project_meta.get("link", "-")),
        ]
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        _set_table_full_width(tbl)

        for label, value in fields:
            row = tbl.add_row()
            lc = row.cells[0]
            lc.text = label
            _set_cell_bg(lc, _C_ORANGE_L)
            _set_cell_padding(lc, 5)
            _set_cell_width(lc, 2.0)
            for r in lc.paragraphs[0].runs:
                r.bold = True
                r.font.size = Pt(10)
            vc = row.cells[1]
            vc.text = _s(value)
            _set_cell_padding(vc, 5)
            for r in vc.paragraphs[0].runs:
                r.font.size = Pt(10)

    except Exception:
        logger.exception("_build_advertiser_profile failed")
        doc.add_paragraph("[광고주 정보 생성 오류]")


def _render_md_body(doc: Document, text: str) -> None:
    """마크다운 텍스트를 Normal/bullet/bold 단락으로 렌더링 (기획 요약 등 본문용)."""
    for line in text.split("\n"):
        s = line.rstrip()
        if not s:
            continue
        if s.startswith("### "):
            p = doc.add_paragraph(s[4:])
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(s[2:], style="List Bullet")
            for r in p.runs:
                r.font.size = Pt(10.5)
        elif s.startswith("---"):
            _add_divider(doc, color="DDDDDD")
        elif s.startswith("|"):
            p = doc.add_paragraph(s)
            for r in p.runs:
                r.font.size = Pt(9)
        else:
            p = doc.add_paragraph(s)
            for r in p.runs:
                r.font.size = Pt(10.5)


def _build_planning_body(doc: Document, ai_content: str) -> None:
    """AI 기획 콘텐츠 — ## 기준 분리, 각 섹션을 Heading 2 + 본문으로 렌더링."""
    try:
        if not ai_content.strip():
            doc.add_paragraph("AI 기획 콘텐츠가 없습니다.")
            return
        parts = re.split(r"(?m)^## ", ai_content)
        for part in parts:
            if not part.strip():
                continue
            first_line, _, body = part.partition("\n")
            header = first_line.strip()
            body   = body.strip()
            if not header:
                continue

            doc.add_heading(header, level=2)
            lower = header.lower()

            if "소식글" in lower:
                for line in body.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            elif "카피" in lower:
                for line in body.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())
            else:
                _render_md_body(doc, body)

            doc.add_paragraph()

    except Exception:
        logger.exception("_build_planning_body failed")
        doc.add_paragraph("[기획 내용 생성 오류]")


# ── Public API ────────────────────────────────────────────────────────────────

def make_charts(
    timeseries: List[TimeseriesRow],
    output_dir: Path,
    mode: str = "db_funnel",
) -> List[Path]:
    """PNG 차트 3개 생성 → 경로 목록 반환.

    출력 파일:
        chart_spend_action.png  — 기간별 광고비용 vs 행동 지표 (dual-axis)
        chart_cpa.png           — CPA / CPE 비용 효율 추이 (line + avg)
        chart_funnel.png        — 전환 퍼널 (수평 막대)

    tracking_mode:
        "db_funnel" (default) : 행동=클릭·문의,  CPA=문의당,  퍼널=노출→클릭→문의→단골
        "landing"             : 행동=클릭·전환,  CPA=전환당,  퍼널=노출→클릭→전환
        "reaction"            : 행동=반응,        CPE=반응당,  퍼널=노출→반응 분해
    """
    if not timeseries:
        return []

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    labels = [_s(r.get("date"), f"기간{i+1}") for i, r in enumerate(timeseries)]
    spends = [_f(r.get("spend")) for r in timeseries]

    charts: List[Path] = []
    for fn, args in [
        (_make_chart_spend_action, (labels, spends, timeseries, mode, output_dir)),
        (_make_chart_cpa_trend,    (labels, spends, timeseries, mode, output_dir)),
        (_make_chart_funnel,       (timeseries, mode, output_dir)),
    ]:
        result = fn(*args)
        if result is not None:
            charts.append(result)

    return charts


def build_report_docx(
    project_meta: ProjectMeta,
    kpi: KPI,
    timeseries: List[TimeseriesRow],
    insights: Insights,
    output_path: Path,
    chart_dir: Optional[Path] = None,
    tracking_mode: str = "db_funnel",
) -> Path:
    """성과보고서 DOCX 생성 → output_path 반환.

    레이아웃 확정안 v1.0 — 10개 섹션:
      Sec 0: 표지               ← _build_cover()
      Sec 1: 목차               ← _build_toc()
      Sec 2: 한 페이지 요약     ← _build_summary_card() + _build_meta_box()
                                   + _build_conclusion() + _build_next_actions()
      Sec 3: 캠페인 개요        ← _build_campaign_overview()
      Sec 4: 성과 요약          ← _build_kpi_table()
      Sec 5: 성과 차트          ← _build_charts_section()
      Sec 6: 인사이트           ← _build_insights_section()
      Sec 7: 다음 실험·개선안   ← _build_experiments_table()
      Sec 8: 판단 기준          ← _build_judgment_criteria()
      Sec 9: 부록               ← _build_appendix()

    tracking_mode: "db_funnel" | "landing" | "reaction"
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if chart_dir is None:
        chart_dir = output_path.parent / "_charts"
    chart_paths = make_charts(timeseries, Path(chart_dir), mode=tracking_mode)

    if _TEMPLATE_REPORT.exists():
        doc = _load_template(_TEMPLATE_REPORT)
    else:
        doc = Document()
        _apply_document_defaults(doc)
    _add_footer(doc, f"{_s(project_meta.get('name'), '광고주')}  |  당근마켓 광고 성과 보고서")

    # Normalize insights once
    norm = _normalize_insights(insights)

    # Sec 0: 표지
    _build_cover(doc, project_meta, "당근마켓 광고 성과 보고서", is_planning=False)
    doc.add_page_break()

    # Sec 1: 목차
    _build_toc(doc)
    doc.add_page_break()

    # Sec 2: 한 페이지 요약
    doc.add_heading("한 페이지 요약", level=1)
    doc.add_paragraph()
    _build_summary_card(doc, kpi, mode=tracking_mode)       # (A) 요약 카드
    doc.add_paragraph()
    _build_meta_box(doc, project_meta, kpi)                  # (B) 메타 박스
    doc.add_paragraph()
    _build_conclusion(doc, norm)                             # (C) 결론
    _build_next_actions(doc, norm)                           # (D) Next Actions
    doc.add_page_break()

    # Sec 3: 캠페인 개요
    doc.add_heading("캠페인 개요", level=1)
    doc.add_paragraph()
    _build_campaign_overview(doc, project_meta, tracking_mode)

    # Sec 4: 성과 요약
    doc.add_heading("성과 요약", level=1)
    doc.add_paragraph()
    _build_kpi_table(doc, kpi, mode=tracking_mode)

    # Sec 5: 성과 차트
    doc.add_heading("성과 차트", level=1)
    _build_charts_section(doc, chart_paths)

    # Sec 6: 인사이트
    doc.add_heading("인사이트", level=1)
    _build_insights_section(doc, norm)

    # Sec 7: 다음 실험·개선안
    doc.add_heading("다음 실험 \u00b7 개선안", level=1)
    _build_experiments_table(doc, norm)

    # Sec 8: 판단 기준
    doc.add_heading("판단 기준", level=1)
    _build_judgment_criteria(doc, norm)

    # Sec 9: 부록
    doc.add_heading("부록", level=1)
    _build_appendix(doc, timeseries, mode=tracking_mode)

    doc.save(output_path)
    return output_path


def build_planning_docx(
    project_meta: ProjectMeta,
    ai_content: str,
    output_path: Path,
) -> Path:
    """기획서 DOCX 생성 → output_path 반환.

    Document structure (document_spec.md v1.0):
      표지
      ── page break ──
      광고주 정보 (Heading 1) + 8행 × 2열 표
      ── page break ──
      AI 생성 기획 내용 (Heading 1)
        1. 기획 요약 / 2. 당근 소식글 / 3. 광고 카피 N개
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if _TEMPLATE_PLAN.exists():
        doc = _load_template(_TEMPLATE_PLAN)
    else:
        doc = Document()
        _apply_document_defaults(doc)
    _add_footer(doc, f"{_s(project_meta.get('name'), '광고주')}  |  당근마켓 광고 기획서")

    # 표지
    _build_cover(doc, project_meta, "당근마켓 광고 기획서", is_planning=True)
    doc.add_page_break()

    # 광고주 정보
    _build_advertiser_profile(doc, project_meta)
    doc.add_page_break()

    # AI 생성 기획 내용
    doc.add_heading("AI 생성 기획 내용", level=1)
    _build_planning_body(doc, ai_content)

    doc.save(output_path)
    return output_path
