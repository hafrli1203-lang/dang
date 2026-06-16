"""광고주 전달용 고급 분석 보고서 DOCX 생성기.

입력:
  - project_meta: {name, region, industry, campaign_name, period, author}
  - ages, campaigns, judgments, plan, priority, var_warnings, pair_gaps
    (모두 app.reporting.demographic 자료구조)
  - ai_sections: {summary, status, findings, plan, expected, client_note, raw}
  - output_path: Path
  - chart_dir: Path (차트 PNG 저장 위치)

출력 섹션:
  1. 표지
  2. 한 줄 요약 (AI)
  3. 현황 진단 (AI 텍스트 + 연령별 CPA 차트 + 자동/수동 비교 차트)
  4. 개선점 (AI 텍스트)
  5. 실행 계획 (AI 텍스트 + 우선순위 표 + 재배분 Before/After 차트)
  6. 예상 효과 (AI 텍스트 + 재배분 시뮬 표)
  7. 광고주 전달 멘트
  8. 부록 — 연령/캠페인 원본 표
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path
from typing import Sequence

import matplotlib
if not matplotlib.is_interactive():
    matplotlib.use("Agg", force=False)
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)

_KOREAN_FONT = "Malgun Gothic"


def _fmt_won(v: float | int) -> str:
    return f"{int(round(v)):,}원"


def _setup_korean_font() -> None:
    try:
        fm.findfont(_KOREAN_FONT, fallback_to_default=False)
        plt.rcParams["font.family"] = _KOREAN_FONT
    except Exception:  # noqa: BLE001
        pass
    plt.rcParams["axes.unicode_minus"] = False


def _add_styled_paragraph(doc: Document, text: str, *, bold: bool = False,
                          size: int = 11, color: RGBColor | None = None,
                          align=None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = _KOREAN_FONT
    if color is not None:
        run.font.color.rgb = color


def _add_markdown_block(doc: Document, text: str) -> None:
    """Render AI markdown-ish block: lines starting with - / 1. / **bold**.

    Keeps it simple — no nested lists; just bullets + bold inline.
    """
    if not text or not text.strip():
        return
    for raw_line in text.split("\n"):
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.lstrip().startswith(("- ", "• ", "* ")):
            _add_bullet(doc, line.lstrip()[2:].strip())
        elif line.lstrip()[:3] in {"1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. ", "8. ", "9. "}:
            _add_number(doc, line.lstrip()[3:].strip())
        else:
            _add_inline(doc, line.strip())


def _add_inline(doc: Document, text: str) -> None:
    """Add a paragraph with **bold** segments split out."""
    p = doc.add_paragraph()
    _emit_runs(p, text)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    _emit_runs(p, text)


def _add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    _emit_runs(p, text)


def _emit_runs(p, text: str) -> None:
    import re
    cursor = 0
    for m in re.finditer(r"\*\*(.+?)\*\*", text):
        if m.start() > cursor:
            r = p.add_run(text[cursor:m.start()])
            r.font.name = _KOREAN_FONT
        r = p.add_run(m.group(1))
        r.bold = True
        r.font.name = _KOREAN_FONT
        cursor = m.end()
    if cursor < len(text):
        r = p.add_run(text[cursor:])
        r.font.name = _KOREAN_FONT


def _make_age_cpa_chart(ages: Sequence, chart_dir: Path) -> Path | None:
    active = [a for a in ages if a.actions > 0]
    if not active:
        return None
    chart_dir.mkdir(parents=True, exist_ok=True)
    _setup_korean_font()

    labels = [a.label for a in active]
    cpas = [a.cpa for a in active]
    avg = sum(cpas) / len(cpas)

    colors = ["#2E7D32" if c <= avg else "#C62828" for c in cpas]

    fig, ax = plt.subplots(figsize=(8, 4), dpi=160)
    bars = ax.bar(labels, cpas, color=colors, edgecolor="#222")
    ax.axhline(avg, linestyle="--", color="#FF6F0F", linewidth=1.5,
               label=f"평균 CPA {int(avg):,}원")
    for b, c in zip(bars, cpas):
        ax.text(b.get_x() + b.get_width() / 2, b.get_height(),
                f"{int(c):,}", ha="center", va="bottom", fontsize=9)
    ax.set_title("연령대별 행동당 비용 (CPA)")
    ax.set_ylabel("CPA (원)")
    ax.legend(loc="upper right")
    ax.grid(axis="y", linestyle=":", alpha=0.4)
    fig.tight_layout()

    out = chart_dir / "analysis_age_cpa.png"
    fig.savefig(out, dpi=160, facecolor="white")
    plt.close(fig)
    return out


def _make_bid_mode_chart(campaigns: Sequence, chart_dir: Path) -> Path | None:
    auto = [c for c in campaigns if c.bid_mode == "auto"]
    manual = [c for c in campaigns if c.bid_mode == "manual"]
    if not (auto and manual):
        return None
    chart_dir.mkdir(parents=True, exist_ok=True)
    _setup_korean_font()

    def _avg_cpa(items: Sequence) -> float:
        cost = sum(c.cost for c in items)
        acts = sum(c.actions for c in items)
        return cost / acts if acts else 0.0

    def _avg_ctr(items: Sequence) -> float:
        return sum(c.ctr for c in items) / len(items) if items else 0.0

    cats = ["수동", "자동"]
    cpas = [_avg_cpa(manual), _avg_cpa(auto)]
    ctrs = [_avg_ctr(manual), _avg_ctr(auto)]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(8, 3.8), dpi=160)
    ax1.bar(cats, cpas, color=["#1565C0", "#FF8F00"], edgecolor="#222")
    for i, v in enumerate(cpas):
        ax1.text(i, v, f"{int(v):,}원", ha="center", va="bottom", fontsize=10)
    ax1.set_title("입찰 모드별 평균 CPA")
    ax1.set_ylabel("CPA (원)")
    ax1.grid(axis="y", linestyle=":", alpha=0.4)

    ax2.bar(cats, ctrs, color=["#1565C0", "#FF8F00"], edgecolor="#222")
    for i, v in enumerate(ctrs):
        ax2.text(i, v, f"{v:.2f}%", ha="center", va="bottom", fontsize=10)
    ax2.set_title("입찰 모드별 평균 CTR")
    ax2.set_ylabel("CTR (%)")
    ax2.grid(axis="y", linestyle=":", alpha=0.4)

    fig.tight_layout()
    out = chart_dir / "analysis_bid_mode.png"
    fig.savefig(out, dpi=160, facecolor="white")
    plt.close(fig)
    return out


def _make_realloc_chart(plan, chart_dir: Path) -> Path | None:
    if plan.current_total <= 0:
        return None
    chart_dir.mkdir(parents=True, exist_ok=True)
    _setup_korean_font()

    cats = ["현재", "조정 후"]
    values = [plan.current_total, plan.projected_total]

    fig, ax = plt.subplots(figsize=(6, 3.6), dpi=160)
    bars = ax.bar(cats, values, color=["#9E9E9E", "#2E7D32"], edgecolor="#222")
    for b, v in zip(bars, values):
        ax.text(b.get_x() + b.get_width() / 2, v,
                f"{v:,}원", ha="center", va="bottom", fontsize=10)
    if plan.savings:
        ax.set_title(
            f"예산 재배분: 절감 {plan.savings:,}원 / 예상 추가 행동 +{plan.expected_action_delta}건"
        )
    else:
        ax.set_title("예산 재배분 시뮬레이션")
    ax.set_ylabel("총예산 (원)")
    ax.grid(axis="y", linestyle=":", alpha=0.4)
    fig.tight_layout()

    out = chart_dir / "analysis_realloc.png"
    fig.savefig(out, dpi=160, facecolor="white")
    plt.close(fig)
    return out


def _build_cover(doc: Document, project_meta: dict) -> None:
    name = project_meta.get("name") or "광고주"
    campaign = project_meta.get("campaign_name") or ""
    region = project_meta.get("region") or ""
    period = project_meta.get("period") or ""
    author = project_meta.get("author") or "당근 광고 도우미"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n광고 성과 분석 보고서")
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.name = _KOREAN_FONT
    run.font.color.rgb = RGBColor(0xFF, 0x6F, 0x0F)

    _add_styled_paragraph(doc, "— 당근 광고 연령/캠페인 최적화 분석 —",
                          size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 2"
    rows = [
        ("광고주", name),
        ("캠페인", campaign or "-"),
        ("지역", region or "-"),
        ("기간", period or "-"),
    ]
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = k
        c1.text = v

    doc.add_paragraph()
    today = datetime.now().strftime("%Y년 %m월 %d일")
    _add_styled_paragraph(doc, f"작성일: {today}",
                          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_styled_paragraph(doc, f"작성: {author}",
                          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def _build_section(doc: Document, heading: str, body: str) -> None:
    doc.add_heading(heading, level=1)
    _add_markdown_block(doc, body or "(데이터 부족)")


def _build_age_table(doc: Document, ages: Sequence) -> None:
    if not ages:
        return
    doc.add_heading("연령대별 성과", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 2"
    hdr = table.rows[0].cells
    for i, h in enumerate(["연령", "비용", "행동", "노출", "클릭", "CPA"]):
        hdr[i].text = h
    for a in ages:
        row = table.add_row().cells
        row[0].text = a.label
        row[1].text = f"{a.cost:,}원"
        row[2].text = f"{a.actions:,}"
        row[3].text = f"{a.impressions:,}"
        row[4].text = f"{a.clicks:,}"
        row[5].text = f"{int(a.cpa):,}원" if a.actions else "-"


def _build_campaign_table(doc: Document, judgments: Sequence) -> None:
    if not judgments:
        return
    doc.add_heading("캠페인별 판정", level=2)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid Accent 2"
    hdr = table.rows[0].cells
    for i, h in enumerate(["캠페인", "입찰", "비용", "행동", "CPA", "판정"]):
        hdr[i].text = h
    for j in judgments:
        row = table.add_row().cells
        row[0].text = j.campaign.name
        row[1].text = {"manual": "수동", "auto": "자동"}.get(j.campaign.bid_mode, "-")
        row[2].text = f"{j.campaign.cost:,}원"
        row[3].text = f"{j.campaign.actions:,}"
        row[4].text = f"{int(j.campaign.cpa):,}원" if j.campaign.actions else "-"
        row[5].text = j.verdict


def _build_funnel_section(doc: Document, funnel, has_conversion_data: bool = True) -> None:
    doc.add_heading("퍼널 단계별 전환", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 2"
    hdr = table.rows[0].cells
    for i, h in enumerate(["단계", "수치", "전환율", "이탈"]):
        hdr[i].text = h
    rows = [
        ("노출", f"{funnel.impressions:,}", "100%", "-"),
        ("클릭", f"{funnel.clicks:,}", f"{funnel.ctr:.2f}% (CTR)",
         f"{funnel.drop_impression_to_click:.1f}%"),
    ]
    # 전환(문의·단골·쿠폰) 데이터가 실제로 수집된 경우에만 행동 단계를 표기한다.
    if has_conversion_data:
        rows.append(
            ("행동", f"{funnel.actions:,}", f"{funnel.cvr:.2f}% (CVR)",
             f"{funnel.drop_click_to_action:.1f}%")
        )
    for r in rows:
        row = table.add_row().cells
        for i, v in enumerate(r):
            row[i].text = v

    if not has_conversion_data:
        _add_styled_paragraph(
            doc,
            "이 파일에는 클릭 이후 전환(문의·단골·쿠폰) 데이터가 수집되지 않아 "
            "노출·클릭까지만 표기했습니다. ‘행동 0’은 성과가 0이라는 뜻이 아니라 "
            "측정값이 없다는 의미이며, 당근 내보내기에 전환 항목을 추가하면 클릭 이후 "
            "단계까지 분석할 수 있습니다.",
            size=10, bold=False, color=RGBColor(0x4B, 0x55, 0x63),
        )
        return
    if funnel.bottleneck:
        drop = (funnel.drop_impression_to_click
                if funnel.bottleneck == "노출→클릭"
                else funnel.drop_click_to_action)
        _add_styled_paragraph(
            doc,
            f"최대 이탈 구간: {funnel.bottleneck} ({drop:.1f}% 이탈) — "
            f"이 구간 개선이 가장 큰 효과를 냅니다.",
            size=11, bold=True, color=RGBColor(0xC6, 0x28, 0x28),
        )


def _build_economics_section(doc: Document, economics) -> None:
    doc.add_heading("MAX CPA · 손익 분석", level=2)
    status_text = {
        "profit": "흑자 / 확장 여력",
        "breakeven": "손익분기 근접",
        "loss": "적자 / 즉시 조정 필요",
        "unknown": "판단 보류",
    }.get(economics.status, economics.status)
    color = {
        "profit": RGBColor(0x2E, 0x7D, 0x32),
        "breakeven": RGBColor(0xFF, 0x8F, 0x00),
        "loss": RGBColor(0xC6, 0x28, 0x28),
    }.get(economics.status, RGBColor(0x42, 0x42, 0x42))

    _add_styled_paragraph(
        doc,
        f"객단가 {economics.avg_order_value:,}원 · "
        f"목표 이익률 {economics.target_margin_rate*100:.0f}% 기준",
        size=10, color=RGBColor(0x80, 0x80, 0x80),
    )

    rows = [
        ("현재 CPA", _fmt_won(economics.current_cpa)),
        ("손익분기 CPA", _fmt_won(economics.breakeven_cpa)),
        ("MAX CPA (목표 이익 반영)", _fmt_won(economics.max_cpa)),
        ("한계 소진율", f"{economics.burn_rate*100:.1f}%"),
        ("판정", status_text),
        ("예상 매출", _fmt_won(economics.expected_revenue)),
        ("광고 후 이익(추정)", _fmt_won(economics.expected_profit)),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 2"
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = k
        c1.text = v

    interpretations = {
        "profit": "현재 CPA가 MAX CPA 대비 안정 구간입니다. 효율 좋은 연령·소재에 예산을 확대해도 됩니다.",
        "breakeven": "현재 CPA가 MAX CPA에 가까워졌습니다. 효율 낮은 구간을 끄고 예산을 재배분하세요.",
        "loss": f"현재 CPA가 MAX CPA를 초과해 적자입니다. {_fmt_won(economics.max_cpa)} 이하로 낮춰야 손익분기. 비효율 연령·캠페인 OFF 즉시 필요.",
        "unknown": "객단가 또는 이익률 정보가 부족합니다.",
    }
    _add_styled_paragraph(
        doc,
        interpretations.get(economics.status, ""),
        size=11, bold=True, color=color,
    )


def _build_realloc_table(doc: Document, plan) -> None:
    doc.add_heading("예산 재배분 시뮬레이션", level=2)
    rows = [
        ("현재 총예산", f"{plan.current_total:,}원"),
        ("절감 예상", f"{plan.savings:,}원"),
        ("추가 행동 예상", f"+{plan.expected_action_delta}건"),
        ("조정 후 총예산", f"{plan.projected_total:,}원"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Light Grid Accent 2"
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = k
        c1.text = v

    if plan.cuts:
        _add_styled_paragraph(doc, "축소 / OFF 대상", bold=True, size=11)
        for n, a in plan.cuts:
            _add_bullet(doc, f"{n}: -{a:,}원")
    if plan.boosts:
        _add_styled_paragraph(doc, "증액 대상", bold=True, size=11)
        for n, a in plan.boosts:
            _add_bullet(doc, f"{n}: +{a:,}원")


def build_analysis_docx(
    project_meta: dict,
    ages: Sequence,
    campaigns: Sequence,
    judgments: Sequence,
    plan,
    priority: Sequence[str],
    var_warnings: Sequence,
    pair_gaps: Sequence,
    ai_sections: dict,
    output_path: Path,
    chart_dir: Path,
    funnel=None,
    economics=None,
    metrics_available=None,
) -> Path:
    """Generate the 광고주 전달용 분석 보고서 DOCX. Returns output_path."""
    doc = Document()
    if metrics_available is not None:
        _avail = set(metrics_available)
        has_conversion_data = any(
            k in _avail for k in ("inquiries", "regulars", "coupons", "actions")
        )
    else:
        has_conversion_data = bool(funnel is not None and funnel.actions > 0)

    _build_cover(doc, project_meta)

    # 1. 한 줄 요약
    doc.add_heading("핵심 요약", level=1)
    summary = ai_sections.get("summary") or _fallback_summary(plan, ages)
    _add_styled_paragraph(doc, summary, size=12, bold=True,
                          color=RGBColor(0xFF, 0x6F, 0x0F))

    # 2. 현황 진단
    doc.add_heading("현황 진단", level=1)
    _add_markdown_block(doc, ai_sections.get("status") or _fallback_status(ages, campaigns, plan))

    # Funnel section
    if funnel is not None and funnel.impressions > 0:
        _build_funnel_section(doc, funnel, has_conversion_data)

    # Economics section
    if economics is not None and economics.avg_order_value > 0:
        _build_economics_section(doc, economics)

    age_chart = _make_age_cpa_chart(ages, chart_dir)
    if age_chart and age_chart.exists():
        doc.add_picture(str(age_chart), width=Cm(15))

    bid_chart = _make_bid_mode_chart(campaigns, chart_dir)
    if bid_chart and bid_chart.exists():
        doc.add_picture(str(bid_chart), width=Cm(15))

    _build_age_table(doc, ages)
    _build_campaign_table(doc, judgments)

    # 3. 개선점
    doc.add_heading("개선점", level=1)
    _add_markdown_block(doc, ai_sections.get("findings") or _fallback_findings(judgments, var_warnings, pair_gaps))

    # 캠페인 비교 경고는 출력하지 않습니다 (사용자 피드백: 비교가 아닌 개별 분석).

    # 4. 실행 계획
    doc.add_heading("실행 계획", level=1)
    _add_markdown_block(doc, ai_sections.get("plan") or _fallback_plan(priority))

    if priority:
        doc.add_heading("우선순위 체크리스트", level=2)
        for item in priority:
            _add_number(doc, item)

    realloc_chart = _make_realloc_chart(plan, chart_dir)
    if realloc_chart and realloc_chart.exists():
        doc.add_picture(str(realloc_chart), width=Cm(13))

    # 5. 예상 효과
    doc.add_heading("예상 효과", level=1)
    _add_markdown_block(doc, ai_sections.get("expected") or _fallback_expected(plan))
    _build_realloc_table(doc, plan)

    # 6. 광고주 전달 멘트
    note = ai_sections.get("client_note", "").strip()
    if note:
        doc.add_heading("광고주 전달 멘트", level=1)
        _add_markdown_block(doc, note)

    # 7. 소식글 카피 수정안 (제목 + 본문 완성형)
    copy_rev = ai_sections.get("copy_revisions", "").strip()
    if copy_rev:
        doc.add_heading("소식글 카피 수정안 (제목 + 본문 완성형)", level=1)
        _add_markdown_block(doc, copy_rev)

    # 8. 썸네일 이미지 생성 프롬프트
    thumb_prompts = ai_sections.get("thumbnail_prompts", "").strip()
    if thumb_prompts:
        doc.add_heading("썸네일 이미지 생성 프롬프트", level=1)
        _add_markdown_block(doc, thumb_prompts)
        _add_styled_paragraph(
            doc,
            "각 프롬프트 블록을 복사해 Gemini / Nano Banana / 기타 이미지 생성기에 "
            "그대로 붙여넣어 새 썸네일을 만드세요.",
            size=9, color=RGBColor(0x80, 0x80, 0x80),
        )

    # Footer
    doc.add_paragraph()
    _add_styled_paragraph(
        doc,
        "본 보고서는 당근 광고 도우미가 자동 생성한 분석 결과입니다. "
        "실제 광고 운영은 데이터 추이를 추가 모니터링하며 조정하세요.",
        size=9, color=RGBColor(0x80, 0x80, 0x80),
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ───────────────── fallback summaries (used when AI block missing) ─────────────────

def _fallback_summary(plan, ages: Sequence) -> str:
    if not ages:
        return "데이터가 부족하여 요약을 생성할 수 없습니다."
    total_actions = sum(a.actions for a in ages)
    total_cost = sum(a.cost for a in ages)
    avg_cpa = total_cost / total_actions if total_actions else 0
    return (
        f"총 {total_cost:,}원으로 {total_actions:,}건의 행동, "
        f"평균 CPA {int(avg_cpa):,}원. "
        f"재배분 시 예상 추가 +{plan.expected_action_delta}건."
    )


def _fallback_status(ages: Sequence, campaigns: Sequence, plan) -> str:
    if not ages and not campaigns:
        return "분석 가능한 breakdown 데이터가 없습니다."
    lines = []
    if ages:
        active = [a for a in ages if a.actions > 0]
        if active:
            best = min(active, key=lambda a: a.cpa)
            worst = max(active, key=lambda a: a.cpa)
            lines.append(f"- 최고 효율 연령: **{best.label}** (CPA {int(best.cpa):,}원)")
            lines.append(f"- 최악 효율 연령: **{worst.label}** (CPA {int(worst.cpa):,}원)")
    if campaigns:
        lines.append(f"- 운영 캠페인 수: {len(campaigns)}개")
    if plan.current_total:
        lines.append(f"- 현재 총예산: {plan.current_total:,}원")
    return "\n".join(lines)


def _fallback_findings(judgments: Sequence, var_warnings: Sequence, pair_gaps: Sequence) -> str:
    lines = []
    for j in judgments:
        if j.verdict in ("캠페인OFF", "소재전면교체", "증액"):
            lines.append(f"- **[{j.verdict}] {j.campaign.name}** — {j.reason}")
    # Note: pair_gaps는 정상 운영 패턴일 수 있어 액션화하지 않음 (단계별 자동/수동 활용)
    return "\n".join(lines) if lines else "현재 데이터 범위에서 즉시 조치할 항목은 없습니다."


def _fallback_plan(priority: Sequence[str]) -> str:
    if not priority:
        return "데이터가 부족합니다."
    return "\n".join(f"{i+1}. {p.split('—', 1)[-1].strip() if '—' in p else p}"
                    for i, p in enumerate(priority))


def _fallback_expected(plan) -> str:
    return (
        f"- 절감 예상: {plan.savings:,}원\n"
        f"- 추가 행동 예상: +{plan.expected_action_delta}건\n"
        f"- 조정 후 총예산: {plan.projected_total:,}원"
    )
