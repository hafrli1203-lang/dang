# -*- coding: utf-8 -*-
"""실측 당근 양식(밀양: 전환 컬럼 없음) 수용 + '있는 것만 솔직히' 보고서 회귀 테스트.

세 가지를 검증한다:
1. 성과보고서가 long-format xlsx를 일자별 rows로 롤업하고, 7열 템플릿은 폴백.
2. build_analysis_prompt가 전환 데이터 없을 때 '100% 손실' 단정을 금지시킨다.
3. build_analysis_docx 퍼널 섹션이 전환 데이터 없을 때 행동 행을 빼고 안내문을 단다.
"""
import io
import tempfile
import unittest
from pathlib import Path

import openpyxl
from docx import Document

from app.pages.report import _rows_from_daangn_breakdown
from app.ai_engine import build_analysis_prompt
from app.reporting.analysis_docx import build_analysis_docx
from app.reporting.demographic import (
    Funnel,
    parse_demographic_xlsx,
    judge_campaigns,
    simulate_reallocation,
)


def _long_format_no_conversion() -> bytes:
    """밀양 양식: 기간×캠페인, 연령·전환 컬럼 없음."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append([
        "기간", "캠페인 이름", "캠페인 ID",
        "비용 (VAT 포함)", "노출 수", "도달 수", "클릭 수",
        "클릭률", "클릭당 비용(CPC)", "노출당 비용(CPM)",
    ])
    ws.append(("2026.06.01.", "밀양_수동_35~59", 1, 8019, 2545, 1645, 25, 0.98, 320, 3150))
    ws.append(("2026.06.02.", "밀양_수동_35~59", 1, 8972, 2998, 2025, 28, 0.93, 320, 2992))
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _seven_col_template() -> bytes:
    """기존 7열 템플릿 (기간|비용|노출|클릭|문의|단골|쿠폰)."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["기간", "비용", "노출", "클릭", "문의", "단골", "쿠폰"])
    ws.append(["1주차", 75000, 12000, 480, 18, 3, 5])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


class TestReportIngestion(unittest.TestCase):
    def test_long_format_rolled_up_by_date(self):
        out = _rows_from_daangn_breakdown(_long_format_no_conversion())
        self.assertIsNotNone(out)
        rows, info = out
        self.assertEqual(len(rows), 2)  # 2 dates
        self.assertEqual(rows[0]["cost"], 8019)
        self.assertEqual(rows[0]["impressions"], 2545)
        # 전환 컬럼이 없으므로 0이되, available에는 들어가지 않음(있는 것만)
        self.assertEqual(rows[0]["inquiries"], 0)
        self.assertEqual(set(info["metrics_available"]), {"impressions", "clicks"})

    def test_seven_col_template_falls_back(self):
        # 7열 템플릿은 timeseries가 없으므로 None → _parse_excel 폴백 경로
        self.assertIsNone(_rows_from_daangn_breakdown(_seven_col_template()))


class TestAnalysisPromptHonesty(unittest.TestCase):
    def _campaigns(self):
        return parse_demographic_xlsx(_long_format_no_conversion())["campaigns"]

    def test_no_conversion_forbids_loss_framing(self):
        camps = self._campaigns()
        funnel = Funnel(impressions=47403, clicks=365, actions=0)
        prompt = build_analysis_prompt(
            project={"name": "지니스안경"}, ages=[], campaigns=camps,
            judgments=[], plan=simulate_reallocation([]), priority=[],
            var_warnings=[], pair_gaps=[], funnel=funnel,
            metrics_available={"impressions", "clicks"},
        )
        self.assertIn("클릭까지만 측정됨", prompt)
        self.assertIn("수집되지 않았습니다", prompt)
        # 측정 안 된 전환을 사실처럼 CVR 수치로 적는 데이터 라인이 없어야 함.
        # (금지 지시문 안의 "'CVR 0%'" 언급은 허용 — 그건 단정을 막는 문구이다.)
        self.assertNotIn("(CVR", prompt)
        self.assertNotIn("→ 행동", prompt)

    def test_with_conversion_keeps_cvr_line(self):
        funnel = Funnel(impressions=14745, clicks=76, actions=17)
        prompt = build_analysis_prompt(
            project={"name": "지니스안경"}, ages=[], campaigns=[],
            judgments=[], plan=simulate_reallocation([]), priority=[],
            var_warnings=[], pair_gaps=[], funnel=funnel,
            metrics_available={"impressions", "clicks", "inquiries", "actions"},
        )
        self.assertIn("CVR", prompt)
        self.assertIn("행동 17", prompt)


class TestAnalysisDocxHonesty(unittest.TestCase):
    def _build(self, funnel, metrics_available):
        with tempfile.TemporaryDirectory() as tmp:
            out = build_analysis_docx(
                project_meta={"name": "지니스안경"},
                ages=[], campaigns=[], judgments=[],
                plan=simulate_reallocation([]), priority=[],
                var_warnings=[], pair_gaps=[],
                ai_sections={"summary": "요약"},
                output_path=Path(tmp) / "a.docx",
                chart_dir=Path(tmp),
                funnel=funnel,
                metrics_available=metrics_available,
            )
            return Document(str(out))

    def test_no_conversion_funnel_omits_action_row(self):
        doc = self._build(Funnel(impressions=47403, clicks=365, actions=0),
                          {"impressions", "clicks"})
        full = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("수집되지 않아", full)
        # 퍼널 표에 '행동' 행이 없어야 함
        cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        self.assertNotIn("행동", cells)

    def test_with_conversion_funnel_has_action_row(self):
        doc = self._build(Funnel(impressions=14745, clicks=76, actions=17),
                          {"impressions", "clicks", "inquiries", "actions"})
        cells = [c.text for t in doc.tables for r in t.rows for c in r.cells]
        self.assertIn("행동", cells)


if __name__ == "__main__":
    unittest.main()
