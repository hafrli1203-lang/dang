"""4-step wizard UI for the planning page news-post tab.

Steps 1-2 are active; Steps 3-4 show placeholders.
Follows the closure-based pattern from proposal_tab.py.
"""
import asyncio
import base64
import json
import logging
import os
import re
from pathlib import Path

from nicegui import ui, app as nicegui_app

from app.theme import section_header
from app.ai_engine import (
    build_strategy_prompt,
    build_planning_prompt,
    build_ad_settings_prompt,
    build_wizard_proposal_prompt,
    parse_strategy_sections,
    parse_ad_settings_sections,
    parse_wizard_proposal_sections,
    SYSTEM_GUIDE_PLANNING,
    SYSTEM_GUIDE_STRATEGY,
    CATEGORIES,
    _AD_SETTINGS_SECTION_NAMES,
    _AD_SETTINGS_SECTION_KEYS,
    _WIZARD_PROPOSAL_SECTION_NAMES,
    _WIZARD_PROPOSAL_SECTION_KEYS,
)
from app.ai.news_post_guard import _split_blocks
from app.ai.providers import get_provider, ClaudeProvider, OpenAIProvider
from app.ai.output_validator import repair_output, get_schema
from app.reporting.docx_report import build_planning_docx, build_proposal_docx
from app.export_manager import ExportManager
from app.database import (
    get_project,
    get_latest_content,
    save_generated_content,
    get_setting,
    save_setting,
    delete_setting,
)
from app.logger import get_logger

_log = get_logger("planning_wizard")

_STEP_LABELS = ["전략 분석", "콘텐츠 생성", "광고 세팅", "운영 제안서"]
_STEP_ICONS = ["analytics", "edit_note", "ads_click", "description"]
_MAX_ENABLED_STEP = 4

_STRATEGY_SECTION_LABELS = {
    "target": "타겟 분석",
    "competition": "경쟁 환경 분석",
    "direction": "전략 방향",
    "campaign_group": "캠페인 그룹 구성",
}
_STRATEGY_SECTION_KEYS = ["target", "competition", "direction", "campaign_group"]


def _parse_planning_sections(content: str) -> dict:
    """Parse planning content into 7 sections (mirrors planning.py logic)."""
    blocks = _split_blocks(content)

    v1_text = ""
    v2_text = ""
    if "의심해소" in blocks:
        v1_text = "【소식글 1 | 의심해소형】\n" + blocks["의심해소"]
    if "가성비" in blocks:
        v2_text = "【소식글 2 | 가성비형】\n" + blocks["가성비"]

    # Fallback: ### 2-A / ### 2-B 또는 ## 2. 소식글 내부 파싱
    if not v1_text or not v2_text:
        # ### 기반 분리 시도
        sub_parts = re.split(r"(?m)^(###\s+2-[A-Ba-b]\.?\s*.+)", content)
        for i, sp in enumerate(sub_parts):
            if not re.match(r"###\s+2-[A-Ba-b]", sp):
                continue
            body = sub_parts[i + 1] if i + 1 < len(sub_parts) else ""
            if ("의심해소" in sp or "2-A" in sp.upper() or "2-a" in sp) and not v1_text:
                v1_text = "【소식글 1 | 의심해소형】\n" + body.strip()
            elif ("가성비" in sp or "2-B" in sp.upper() or "2-b" in sp) and not v2_text:
                v2_text = "【소식글 2 | 가성비형】\n" + body.strip()

    # Fallback 2: ## 2. 소식글 섹션 내에서 "의심해소"/"가성비" 키워드로 분리
    if not v1_text or not v2_text:
        sec2_match = re.search(r"(?m)^## 2\.\s*.+?\n(.*?)(?=^## [3-9]\.|\Z)", content, re.DOTALL)
        if sec2_match:
            sec2_body = sec2_match.group(1)
            # "의심해소" ~ "가성비" 사이를 v1, "가성비" 이후를 v2로 분리
            v_split = re.split(r"(?m)^(.+가성비.+)$", sec2_body, maxsplit=1)
            if len(v_split) >= 3:
                if not v1_text and "의심해소" in v_split[0]:
                    v1_text = "【소식글 1 | 의심해소형】\n" + v_split[0].strip()
                if not v2_text:
                    v2_text = "【소식글 2 | 가성비형】\n" + (v_split[1] + v_split[2]).strip()

    summary = ""
    ad_copies = ""
    campaign_groups = ""
    thumbnail_guide = ""
    coupon_spec = ""
    naming_convention = ""

    # Collect all ## sections with their headers and bodies
    parts = re.split(r"(?m)^(## .+)", content)
    header_body_pairs = []
    for i, part in enumerate(parts):
        if not part.startswith("## "):
            continue
        body = parts[i + 1] if i + 1 < len(parts) else ""
        header_body_pairs.append((part, body))

    for header, body in header_body_pairs:
        h = header
        # summary: 기획+요약 OR 요약 alone
        if ("기획" in h and "요약" in h) or ("요약" in h and not summary):
            summary = body
        # ad_copies: 카피 OR 제목 OR (광고 with digit or standalone)
        elif "카피" in h or "제목" in h or ("광고" in h and (re.search(r"\d", h) or "카피" not in h)):
            if not ad_copies:
                ad_copies = body
        # campaign_groups: 캠페인+그룹/구성 OR 타겟 OR 그룹 OR 연령 OR 나이
        elif ("캠페인" in h and ("그룹" in h or "구성" in h)) or "타겟" in h or "그룹" in h or "연령" in h or "나이" in h:
            if not campaign_groups:
                campaign_groups = body
        # thumbnail_guide: 썸네일 OR 이미지 OR 촬영 OR 사진
        elif "썸네일" in h or "이미지" in h or "촬영" in h or "사진" in h:
            if not thumbnail_guide:
                thumbnail_guide = body
        # coupon_spec: 쿠폰 OR 혜택 OR 할인
        elif "쿠폰" in h or "혜택" in h or "할인" in h:
            if not coupon_spec:
                coupon_spec = body
        # naming_convention: 네이밍 OR 캠페인명 OR 명명 OR 이름
        elif "네이밍" in h or "캠페인명" in h or "명명" in h or "이름" in h:
            if not naming_convention:
                naming_convention = body

    # Position-based fallback: assign remaining unmatched ## sections (after sosigeul sections)
    # Order expected: summary(1), ad_copies(2), campaign_groups(3), thumbnail_guide(4), coupon_spec(5), naming_convention(6)
    _fallback_keys = ["summary", "ad_copies", "campaign_groups", "thumbnail_guide", "coupon_spec", "naming_convention"]
    _fallback_values = [summary, ad_copies, campaign_groups, thumbnail_guide, coupon_spec, naming_convention]
    matched_headers = set()

    # Re-identify which headers were matched
    for header, body in header_body_pairs:
        h = header
        if ("기획" in h and "요약" in h) or ("요약" in h):
            matched_headers.add(header)
        elif "카피" in h or "제목" in h or ("광고" in h and (re.search(r"\d", h) or "카피" not in h)):
            matched_headers.add(header)
        elif ("캠페인" in h and ("그룹" in h or "구성" in h)) or "타겟" in h or "그룹" in h or "연령" in h or "나이" in h:
            matched_headers.add(header)
        elif "썸네일" in h or "이미지" in h or "촬영" in h or "사진" in h:
            matched_headers.add(header)
        elif "쿠폰" in h or "혜택" in h or "할인" in h:
            matched_headers.add(header)
        elif "네이밍" in h or "캠페인명" in h or "명명" in h or "이름" in h:
            matched_headers.add(header)

    unmatched = [(h, b) for h, b in header_body_pairs if h not in matched_headers]

    # Fill empty slots in order from unmatched sections
    empty_slots = [k for k, v in zip(_fallback_keys, _fallback_values) if not v]
    for slot_key, (uh, ub) in zip(empty_slots, unmatched):
        if slot_key == "summary":
            summary = ub
        elif slot_key == "ad_copies":
            ad_copies = ub
        elif slot_key == "campaign_groups":
            campaign_groups = ub
        elif slot_key == "thumbnail_guide":
            thumbnail_guide = ub
        elif slot_key == "coupon_spec":
            coupon_spec = ub
        elif slot_key == "naming_convention":
            naming_convention = ub

    return {
        "version_1": v1_text,
        "version_2": v2_text,
        "summary": summary,
        "ad_copies": ad_copies,
        "campaign_groups": campaign_groups,
        "thumbnail_guide": thumbnail_guide,
        "coupon_spec": coupon_spec,
        "naming_convention": naming_convention,
        "raw": content,
    }


def build_wizard_ui(
    engine_radio,
    category_sel,
    strategy_sel,
    extra_input,
    prompt_editor,
) -> None:
    """Build the 4-step wizard UI inside the news-post tab panel.

    Parameters are NiceGUI widget references from the parent planning page
    so the wizard can read their current values.
    """

    # -- Wizard state --
    _wizard_state: dict = {
        "current_step": 1,
        "step1_content": "",
        "step1_sections": {},
        "step2_content": "",
        "step2_engine": "claude",
        "step3_content": "",
        "step3_sections": {},
        "step4_content": "",
        "step4_sections": {},
        "project_id": nicegui_app.storage.user.get("current_project_id"),
    }

    # ── Step indicator ──────────────────────────────────────────────────────

    step_indicator = ui.element("div").classes("dg-wizard-steps w-full")

    step_circles: list = []
    step_labels_ui: list = []

    def _render_step_indicator() -> None:
        step_indicator.clear()
        step_circles.clear()
        step_labels_ui.clear()
        current = _wizard_state["current_step"]
        with step_indicator:
            for idx in range(4):
                step_num = idx + 1
                if step_num < current:
                    cls = "completed"
                elif step_num == current:
                    cls = "active"
                elif step_num > _MAX_ENABLED_STEP:
                    cls = "disabled"
                else:
                    cls = "pending"

                with ui.element("div").classes(f"dg-wizard-step {cls}"):
                    if idx < 3:
                        ui.element("div").classes("dg-wizard-step-line")
                    circle = ui.element("div").classes("dg-wizard-step-circle")
                    with circle:
                        if step_num < current:
                            ui.icon("check", size="18px")
                        else:
                            ui.label(str(step_num))
                    lbl = ui.label(_STEP_LABELS[idx]).classes("dg-wizard-step-label")
                    step_circles.append(circle)
                    step_labels_ui.append(lbl)

                    if step_num <= _MAX_ENABLED_STEP:
                        circle.on(
                            "click",
                            lambda _e, s=step_num: _go_to_step(s),
                        )

    # ── Content containers ──────────────────────────────────────────────────

    content_container = ui.column().classes("dg-wizard-content w-full gap-4")

    # ── Step navigation ─────────────────────────────────────────────────────

    def _go_to_step(step: int) -> None:
        if step < 1 or step > _MAX_ENABLED_STEP:
            return
        _wizard_state["current_step"] = step
        _render_step_indicator()
        _render_current_step()

    # ── Unified export bar helper ─────────────────────────────────────────

    def _render_export_bar(
        step_num: int,
        *,
        docx_handler=None,
        pdf_handler=None,
    ) -> None:
        """Render a unified export + navigation button bar.

        Parameters
        ----------
        step_num : int
            Current wizard step (1-4).
        docx_handler : callable, optional
            Custom DOCX export handler.  Falls back to ``_export_step_content``.
        pdf_handler : callable, optional
            Custom PDF export handler.  Falls back to ``_export_step_content``.
        """
        with ui.row().classes("w-full justify-center items-center gap-3 my-3"):
            if step_num > 1:
                ui.button(
                    "이전 단계로", icon="arrow_back",
                    on_click=lambda s=step_num: _go_to_step(s - 1),
                ).classes("dg-btn-secondary dg-btn-sm")

            # DOCX
            if docx_handler is not None:
                ui.button(
                    "DOCX 저장", icon="save",
                    on_click=docx_handler,
                ).classes("dg-btn-success dg-btn-sm")
            else:
                ui.button(
                    "DOCX 저장", icon="save",
                    on_click=lambda s=step_num: _export_step_content(s, "docx"),
                ).classes("dg-btn-success dg-btn-sm")

            # PDF
            if pdf_handler is not None:
                ui.button(
                    "PDF 저장", icon="picture_as_pdf",
                    on_click=pdf_handler,
                ).classes("dg-btn-secondary dg-btn-sm")
            else:
                ui.button(
                    "PDF 저장", icon="picture_as_pdf",
                    on_click=lambda s=step_num: _export_step_content(s, "pdf"),
                ).classes("dg-btn-secondary dg-btn-sm")

            if step_num < 4:
                ui.button(
                    "다음 단계로", icon="arrow_forward",
                    on_click=lambda s=step_num: _go_to_step(s + 1),
                ).classes("dg-btn-primary dg-btn-sm")

    # ── Step 1: Strategy Analysis ───────────────────────────────────────────

    def _render_step1() -> None:
        content_container.clear()
        with content_container:
            section_header("analytics", "전략 분석", "AI가 타겟, 경쟁 환경, 전략 방향을 분석해 드려요.")

            # Show existing strategy if loaded
            if _wizard_state["step1_content"]:
                _render_strategy_result()
                return

            with ui.card().classes("dg-card w-full"):
                with ui.column().classes("w-full gap-3 items-center"):
                    ui.icon("psychology", size="48px").style("color: var(--dg-border)")
                    ui.label(
                        "프로젝트 정보를 바탕으로 전략 분석을 만들어 드려요."
                    ).classes("dg-text-sm").style("text-align: center")
                    ui.label(
                        "타겟 페르소나, 경쟁 환경, 전략 방향, 캠페인 그룹 구성이 담겨요."
                    ).classes("dg-label-sm").style("text-align: center")

                    with ui.row().classes("gap-3 items-center"):
                        gen_strategy_btn = ui.button(
                            "전략 분석 생성", icon="auto_awesome",
                            on_click=lambda: _generate_strategy(),
                        ).classes("dg-btn-primary")
                        s1_spinner = ui.spinner(size="28px").classes("hidden")
                        s1_status = ui.label("").classes("dg-progress-text hidden")

                    # Store refs for the generate handler
                    _wizard_state["_s1_btn"] = gen_strategy_btn
                    _wizard_state["_s1_spinner"] = s1_spinner
                    _wizard_state["_s1_status"] = s1_status

    def _render_strategy_result() -> None:
        """Render strategy analysis results with edit/regenerate controls."""
        content_container.clear()
        sections = _wizard_state["step1_sections"]

        with content_container:
            section_header("analytics", "전략 분석 결과")

            _render_export_bar(1)

            for key in _STRATEGY_SECTION_KEYS:
                label = _STRATEGY_SECTION_LABELS[key]
                body = sections.get(key, "(아직 내용이 없어요)")
                idx = _STRATEGY_SECTION_KEYS.index(key) + 1

                with ui.expansion(
                    f"{idx}. {label}",
                    icon="article",
                    value=True,
                ).classes("w-full dg-expansion"):
                    md_widget = ui.markdown(body).classes("w-full dg-prose")
                    edit_area = ui.textarea(value=body).classes(
                        "w-full hidden dg-input"
                    ).props("rows=8 outlined")

                    with ui.row().classes("gap-2 mt-2"):
                        edit_btn = ui.button("편집", icon="edit").classes("dg-btn-ghost dg-btn-sm")
                        save_btn = ui.button("편집 완료", icon="check").classes("dg-btn-success dg-btn-sm hidden")

                        def _toggle_edit(
                            _e, _md=md_widget, _ea=edit_area, _eb=edit_btn, _sb=save_btn,
                        ) -> None:
                            _md.classes("hidden", remove=False)
                            _ea.classes(remove="hidden")
                            _eb.classes("hidden", remove=False)
                            _sb.classes(remove="hidden")

                        def _save_edit(
                            _e, _md=md_widget, _ea=edit_area, _eb=edit_btn, _sb=save_btn,
                            _key=key,
                        ) -> None:
                            new_text = _ea.value
                            _wizard_state["step1_sections"][_key] = new_text
                            _md.set_content(new_text)
                            _md.classes(remove="hidden")
                            _ea.classes("hidden", remove=False)
                            _sb.classes("hidden", remove=False)
                            _eb.classes(remove="hidden")
                            # Rebuild raw content
                            parts = []
                            for i, k in enumerate(_STRATEGY_SECTION_KEYS):
                                n = _STRATEGY_SECTION_LABELS[k]
                                b = _wizard_state["step1_sections"].get(k, "")
                                parts.append(f"## {i + 1}. {n}\n{b}")
                            _wizard_state["step1_content"] = "\n\n".join(parts)
                            # Save to DB
                            pid = nicegui_app.storage.user.get("current_project_id")
                            if pid:
                                save_generated_content(pid, "strategy", _wizard_state["step1_content"], content_type="strategy")
                            ui.notify("수정한 내용을 저장했어요.", type="positive")

                        edit_btn.on_click(_toggle_edit)
                        save_btn.on_click(_save_edit)

            # Feedback + regenerate
            with ui.card().classes("dg-card-flat w-full"):
                ui.label("수정 요청").style("font-weight: 600; font-size: 14px; color: var(--dg-text-primary)")
                feedback_area = ui.textarea(
                    placeholder="예: 20대 타겟을 더 강화해주세요, 긴급성 전략으로 변경 등"
                ).classes("w-full dg-input").props("rows=2 outlined")

                with ui.row().classes("gap-3 items-center"):
                    regen_btn = ui.button(
                        "재생성", icon="refresh",
                        on_click=lambda: _regenerate_strategy(feedback_area.value),
                    ).classes("dg-btn-ghost")
                    regen_spinner = ui.spinner(size="24px").classes("hidden")
                    regen_status = ui.label("").classes("dg-progress-text hidden")

                _wizard_state["_regen_spinner"] = regen_spinner
                _wizard_state["_regen_status"] = regen_status
                _wizard_state["_regen_btn"] = regen_btn

            # Navigation
            _render_export_bar(1)

    async def _generate_strategy() -> None:
        pid = nicegui_app.storage.user.get("current_project_id")
        if not pid:
            ui.notify("프로젝트를 먼저 선택해 주세요.", type="warning")
            return
        project = get_project(pid)
        if not project:
            ui.notify("프로젝트를 찾을 수 없어요. 프로젝트 페이지에서 다시 선택해 주세요.", type="negative")
            return

        engine = engine_radio.value

        # Validate API key
        # Claude는 기본이 CLI(구독). API 백엔드일 때만 키 필요.
        if (engine == "claude" and os.getenv("CLAUDE_BACKEND", "cli").strip().lower() == "api"
                and not os.getenv("ANTHROPIC_API_KEY", "")):
            ui.notify("Claude를 API 모드로 쓰려면 .env에 ANTHROPIC_API_KEY가 필요해요. (CLI 모드는 키 없이 동작)", type="negative")
            return
        # GPT는 기본이 codex CLI(구독). API 백엔드일 때만 키 필요.
        if (engine in ("gpt", "coordinate") and os.getenv("OPENAI_BACKEND", "cli").strip().lower() == "api"
                and not os.getenv("OPENAI_API_KEY", "")):
            ui.notify("GPT를 API 모드로 쓰려면 .env에 OPENAI_API_KEY가 필요해요. (CLI 모드는 codex 로그인으로 동작)", type="negative")
            return

        btn = _wizard_state.get("_s1_btn")
        spinner = _wizard_state.get("_s1_spinner")
        status = _wizard_state.get("_s1_status")

        if btn:
            btn.props("disabled loading")
        if spinner:
            spinner.classes(remove="hidden")
        if status:
            status.classes(remove="hidden")
            status.set_text("전략 분석을 만들고 있어요...")

        try:
            guide, prompt = build_strategy_prompt(project)
            loop = asyncio.get_running_loop()
            if engine == "coordinate":
                from app.ai.coordination import coordinate_generate
                content = await coordinate_generate(
                    loop, prompt, guide, "전략 분석",
                    on_drafts=(lambda: status.set_text("Claude와 GPT가 각자 분석하고 있어요...")) if status else None,
                    on_synth=(lambda: status.set_text("Claude가 두 분석을 종합하고 있어요...")) if status else None,
                )
            else:
                provider = get_provider(engine)
                content = await loop.run_in_executor(
                    None, lambda: provider.generate_text(prompt, system_prompt=guide),
                )

            # 누락/부실 섹션이 있으면 1회 보정 (best-effort, 실패 시 원본 유지).
            if status:
                status.set_text("빠진 부분이 없는지 확인하고 있어요...")
            content = await loop.run_in_executor(
                None, lambda: repair_output(content, get_schema("strategy"), engine=engine),
            )

            _wizard_state["step1_content"] = content
            _wizard_state["step1_sections"] = parse_strategy_sections(content)
            _wizard_state["project_id"] = pid

            # Save to DB
            save_generated_content(pid, engine, content, content_type="strategy")

            _render_strategy_result()
            ui.notify("전략 분석이 완성됐어요!", type="positive")

        except Exception as exc:
            _log.exception("전략 분석 생성 실패: %s", exc)
            ui.notify(f"전략 분석을 만들지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative", timeout=8000)
            if status:
                status.set_text(f"오류: {exc}")
        finally:
            if btn:
                btn.props(remove="disabled loading")
            if spinner:
                spinner.classes("hidden")

    async def _regenerate_strategy(feedback: str = "") -> None:
        pid = nicegui_app.storage.user.get("current_project_id")
        if not pid:
            return
        project = get_project(pid)
        if not project:
            return

        engine = engine_radio.value if engine_radio.value not in ("both", "coordinate") else "claude"

        regen_spinner = _wizard_state.get("_regen_spinner")
        regen_status = _wizard_state.get("_regen_status")
        regen_btn = _wizard_state.get("_regen_btn")

        if regen_btn:
            regen_btn.props("disabled loading")
        if regen_spinner:
            regen_spinner.classes(remove="hidden")
        if regen_status:
            regen_status.classes(remove="hidden")
            regen_status.set_text("다시 만들고 있어요...")

        try:
            guide, prompt = build_strategy_prompt(project)
            if feedback.strip():
                prompt += f"\n\n[수정 요청]\n{feedback.strip()}"
            if _wizard_state["step1_content"]:
                prompt += f"\n\n[이전 분석 결과 (참고하되 수정 요청 반영)]\n{_wizard_state['step1_content']}"

            loop = asyncio.get_running_loop()
            provider = get_provider(engine)
            content = await loop.run_in_executor(
                None, lambda: provider.generate_text(prompt, system_prompt=guide),
            )

            _wizard_state["step1_content"] = content
            _wizard_state["step1_sections"] = parse_strategy_sections(content)

            save_generated_content(pid, engine, content, content_type="strategy")

            _render_strategy_result()
            ui.notify("전략 분석을 다시 만들었어요!", type="positive")

        except Exception as exc:
            _log.exception("전략 분석 재생성 실패: %s", exc)
            ui.notify(f"다시 만들지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative", timeout=8000)
        finally:
            if regen_btn:
                regen_btn.props(remove="disabled loading")
            if regen_spinner:
                regen_spinner.classes("hidden")

    def _advance_to_step2() -> None:
        if not _wizard_state["step1_content"]:
            ui.notify("전략 분석을 먼저 만들어 주세요.", type="warning")
            return
        _wizard_state["current_step"] = 2
        _render_step_indicator()
        _render_current_step()

    # ── Step 2: Content Generation ──────────────────────────────────────────

    def _render_step2() -> None:
        content_container.clear()
        with content_container:
            section_header("edit_note", "콘텐츠 생성", "전략 분석 결과를 반영해 소식글과 기획 콘텐츠를 만들어요.")

            # Strategy summary (collapsible)
            if _wizard_state["step1_content"]:
                with ui.expansion(
                    "전략 분석 요약 (Step 1 결과)", icon="analytics",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step1_content"]).classes("w-full dg-prose")

            # State for step 2
            _s2: dict = {
                "content": _wizard_state.get("step2_content", ""),
                "engine": _wizard_state.get("step2_engine", "claude"),
                "cancelled": False,
            }

            # Action buttons
            with ui.row().classes("gap-3 items-center"):
                gen_btn = ui.button(
                    "기획 콘텐츠 생성", icon="auto_awesome",
                    on_click=lambda: _generate_content(),
                ).classes("dg-btn-primary")
                cancel_btn = ui.button(
                    "중단", icon="stop",
                    on_click=lambda: _s2.__setitem__("cancelled", True),
                ).classes("dg-btn-danger dg-btn-sm hidden")
                spinner = ui.spinner(size="32px").classes("hidden")
                step_label = ui.label("").classes("dg-progress-text hidden")
                download_status = ui.label("").style(
                    "font-size: 13px; font-weight: 600; color: var(--dg-success)"
                ).classes("hidden")

            _render_export_bar(
                2,
                docx_handler=lambda: _export_default(),
                pdf_handler=lambda: _export_default_pdf(),
            )

            def _set_step(text: str) -> None:
                step_label.classes(remove="hidden")
                step_label.set_text(text)

            # Result area
            result_card = ui.card().classes("dg-card w-full hidden")
            with result_card:
                with ui.tabs().classes("w-full dg-tabs") as result_tabs:
                    tab_all = ui.tab("전체 보기")
                    tab_v1 = ui.tab("소식글 1 (의심해소)")
                    tab_v2 = ui.tab("소식글 2 (가성비)")
                    tab_campaign = ui.tab("캠페인 구성")
                    tab_mobile = ui.tab("모바일 미리보기")
                    tab_thumbnail_gen = ui.tab("썸네일 생성")

                with ui.tab_panels(result_tabs, value=tab_all).classes("w-full"):
                    with ui.tab_panel(tab_all):
                        result_md_all = ui.markdown("").classes("w-full dg-prose")
                        edit_area_all = ui.textarea(value="").classes(
                            "w-full hidden dg-input"
                        ).props("rows=20 outlined")
                        with ui.row().classes("gap-2 mt-2"):
                            edit_btn_all = ui.button("편집", icon="edit").classes("dg-btn-ghost dg-btn-sm")
                            save_btn_all = ui.button("편집 완료", icon="check").classes("dg-btn-success dg-btn-sm hidden")
                            cancel_edit_btn_all = ui.button("취소", icon="close").classes("dg-btn-ghost dg-btn-sm hidden")

                        def _toggle_edit_s2(
                            _e, _md=result_md_all, _ea=edit_area_all,
                            _eb=edit_btn_all, _sb=save_btn_all, _cb=cancel_edit_btn_all,
                        ) -> None:
                            _ea.value = _s2.get("content", "")
                            _md.classes("hidden", remove=False)
                            _ea.classes(remove="hidden")
                            _eb.classes("hidden", remove=False)
                            _sb.classes(remove="hidden")
                            _cb.classes(remove="hidden")

                        def _save_edit_s2(
                            _e, _md=result_md_all, _ea=edit_area_all,
                            _eb=edit_btn_all, _sb=save_btn_all, _cb=cancel_edit_btn_all,
                        ) -> None:
                            new_text = _ea.value
                            _s2["content"] = new_text
                            _wizard_state["step2_content"] = new_text
                            _md.classes(remove="hidden")
                            _ea.classes("hidden", remove=False)
                            _sb.classes("hidden", remove=False)
                            _cb.classes("hidden", remove=False)
                            _eb.classes(remove="hidden")
                            pid = nicegui_app.storage.user.get("current_project_id")
                            if pid:
                                save_generated_content(pid, "edited", new_text)
                            _update_result_display(new_text)
                            ui.notify("수정한 내용을 저장했어요.", type="positive")

                        def _cancel_edit_s2(
                            _e, _md=result_md_all, _ea=edit_area_all,
                            _eb=edit_btn_all, _sb=save_btn_all, _cb=cancel_edit_btn_all,
                        ) -> None:
                            _md.classes(remove="hidden")
                            _ea.classes("hidden", remove=False)
                            _sb.classes("hidden", remove=False)
                            _cb.classes("hidden", remove=False)
                            _eb.classes(remove="hidden")

                        edit_btn_all.on_click(_toggle_edit_s2)
                        save_btn_all.on_click(_save_edit_s2)
                        cancel_edit_btn_all.on_click(_cancel_edit_s2)
                    with ui.tab_panel(tab_v1):
                        with ui.row().classes("w-full justify-end mb-1"):
                            ui.button(
                                "복사", icon="content_copy",
                                on_click=lambda: _copy_section("version_1"),
                            ).classes("dg-btn-ghost dg-btn-sm")
                        result_card_v1 = ui.column().classes("w-full gap-3")
                    with ui.tab_panel(tab_v2):
                        with ui.row().classes("w-full justify-end mb-1"):
                            ui.button(
                                "복사", icon="content_copy",
                                on_click=lambda: _copy_section("version_2"),
                            ).classes("dg-btn-ghost dg-btn-sm")
                        result_card_v2 = ui.column().classes("w-full gap-3")
                    with ui.tab_panel(tab_campaign):
                        with ui.column().classes("w-full gap-4"):
                            with ui.element("div").classes("w-full"):
                                ui.label("광고 카피 9종").style(
                                    "font-size: 15px; font-weight: 700; color: var(--dg-text-primary); margin-bottom: 8px;"
                                )
                                result_md_copies = ui.markdown("").classes("w-full dg-prose")
                            with ui.element("div").classes("w-full"):
                                ui.label("캠페인 그룹 구성").style(
                                    "font-size: 15px; font-weight: 700; color: var(--dg-text-primary); margin-bottom: 8px;"
                                )
                                result_md_campaign = ui.markdown("").classes("w-full dg-prose")
                            with ui.row().classes("w-full gap-4 flex-wrap"):
                                with ui.column().classes("flex-1 min-w-[280px]"):
                                    ui.label("썸네일/이미지 가이드").style(
                                        "font-size: 15px; font-weight: 700; color: var(--dg-text-primary); margin-bottom: 8px;"
                                    )
                                    result_md_thumbnail = ui.markdown("").classes("w-full dg-prose")
                                with ui.column().classes("flex-1 min-w-[280px]"):
                                    ui.label("쿠폰 스펙 / 네이밍").style(
                                        "font-size: 15px; font-weight: 700; color: var(--dg-text-primary); margin-bottom: 8px;"
                                    )
                                    result_md_coupon = ui.markdown("").classes("w-full dg-prose")
                    with ui.tab_panel(tab_mobile):
                        with ui.row().classes("w-full justify-center gap-6 flex-wrap"):
                            mobile_frame_v1 = ui.column().classes("items-center gap-2")
                            mobile_frame_v2 = ui.column().classes("items-center gap-2")

                    with ui.tab_panel(tab_thumbnail_gen):
                        with ui.row().classes("w-full gap-4 flex-wrap"):
                            # Left: AI guide
                            with ui.column().classes("flex-1 min-w-[300px] gap-2"):
                                ui.label("AI 썸네일 가이드").style(
                                    "font-size: 15px; font-weight: 700; color: var(--dg-text-primary);"
                                )
                                result_md_thumbnail_gen = ui.markdown(
                                    "*기획 콘텐츠를 먼저 만들면 AI 가이드가 여기에 표시돼요.*"
                                ).classes("w-full dg-prose")

                            # Right: generation UI
                            with ui.column().classes("flex-1 min-w-[300px] gap-3"):
                                ui.label("썸네일 생성").style(
                                    "font-size: 15px; font-weight: 700; color: var(--dg-text-primary);"
                                )

                                # Reference image upload
                                _thumb_state: dict = {
                                    "ref_bytes": None,
                                    "ref_mime": "image/png",
                                    "result_bytes": None,
                                    "result_mime": "image/png",
                                }

                                thumb_ref_preview = ui.column().classes("w-full hidden")

                                with ui.row().classes("gap-3 items-center"):
                                    thumb_ref_upload = ui.upload(
                                        label="참고 이미지 (선택)",
                                        auto_upload=True,
                                        on_upload=lambda e: _handle_thumb_ref_upload(e),
                                    ).classes("dg-upload").props('accept="image/*"')

                                    thumb_ref_clear_btn = ui.button(
                                        "제거", icon="close",
                                        on_click=lambda: _clear_thumb_ref(),
                                    ).classes("dg-btn-secondary dg-btn-sm hidden")

                                # Prompt input
                                thumb_prompt = ui.textarea(
                                    label="생성 프롬프트",
                                    placeholder=(
                                        "예: 당근마켓 스타일의 따뜻한 오렌지톤 썸네일, "
                                        "\"신선한 제철 과일 50% 할인\" 문구 포함"
                                    ),
                                ).classes("w-full dg-input").props("rows=3 outlined")

                                # Action row
                                with ui.row().classes("gap-2 items-center"):
                                    thumb_gen_btn = ui.button(
                                        "썸네일 생성", icon="auto_awesome",
                                        on_click=lambda: _generate_thumbnail(),
                                    ).classes("dg-btn-primary")
                                    thumb_spinner = ui.spinner(size="24px").classes("hidden")
                                    thumb_status = ui.label("").classes("dg-progress-text hidden")

                                # Result preview
                                thumb_result_container = ui.column().classes("w-full items-center gap-2 hidden")
                                with thumb_result_container:
                                    thumb_result_img = ui.image("").classes("w-full max-w-[400px] dg-image-preview")
                                    thumb_result_label = ui.label("").classes("dg-label-sm")

                                # Save button
                                thumb_save_btn = ui.button(
                                    "기본 폴더에 저장", icon="save",
                                    on_click=lambda: _save_thumbnail(),
                                ).classes("dg-btn-success hidden")

            # ── Thumbnail generation helpers ──────────────────────────────────

            async def _handle_thumb_ref_upload(e) -> None:
                try:
                    data = await e.file.read()
                    mime = e.file.content_type or "image/png"
                    _thumb_state["ref_bytes"] = data
                    _thumb_state["ref_mime"] = mime
                    b64 = base64.b64encode(data).decode()
                    thumb_ref_preview.clear()
                    thumb_ref_preview.classes(remove="hidden")
                    with thumb_ref_preview:
                        ui.image(f"data:{mime};base64,{b64}").classes("max-w-xs dg-image-preview")
                        ui.label(f"{len(data):,} bytes").classes("dg-label-sm")
                    thumb_ref_clear_btn.classes(remove="hidden")
                except Exception as exc:
                    ui.notify(f"이미지를 읽지 못했어요. 파일을 확인하고 다시 올려 주세요. ({exc})", type="negative")

            def _clear_thumb_ref() -> None:
                _thumb_state["ref_bytes"] = None
                _thumb_state["ref_mime"] = "image/png"
                thumb_ref_preview.clear()
                thumb_ref_preview.classes("hidden")
                thumb_ref_clear_btn.classes("hidden")
                thumb_ref_upload.reset()

            async def _generate_thumbnail() -> None:
                prompt_text = thumb_prompt.value.strip()
                if not prompt_text:
                    ui.notify("만들고 싶은 썸네일을 먼저 적어 주세요.", type="warning")
                    return

                thumb_gen_btn.props("disabled loading")
                thumb_spinner.classes(remove="hidden")
                thumb_status.classes(remove="hidden")
                thumb_status.set_text("썸네일을 만들고 있어요...")

                try:
                    from app.ai.image_provider import get_image_provider
                    from app.ai.thumbnail_style import build_natural_thumbnail_prompt
                    provider = get_image_provider()
                    loop = asyncio.get_running_loop()
                    ref = _thumb_state["ref_bytes"]
                    ref_mime = _thumb_state["ref_mime"]

                    # 당근 피드용 자연 실사로 강제 — 광고 티가 나면 스크롤로 넘어간다.
                    final_prompt = build_natural_thumbnail_prompt(
                        prompt_text, has_reference=ref is not None
                    )

                    if ref is not None:
                        img_bytes, mime = await loop.run_in_executor(
                            None,
                            lambda: provider.generate_image(
                                final_prompt,
                                reference_image=ref,
                                reference_mime=ref_mime,
                            ),
                        )
                    else:
                        img_bytes, mime = await loop.run_in_executor(
                            None,
                            lambda: provider.generate_image(final_prompt),
                        )

                    _thumb_state["result_bytes"] = img_bytes
                    _thumb_state["result_mime"] = mime

                    b64 = base64.b64encode(img_bytes).decode()
                    thumb_result_img.set_source(f"data:{mime};base64,{b64}")
                    thumb_result_label.set_text(f"{len(img_bytes):,} bytes | {mime}")
                    thumb_result_container.classes(remove="hidden")
                    thumb_save_btn.classes(remove="hidden")
                    thumb_status.set_text("썸네일이 완성됐어요!")
                    ui.notify("썸네일이 완성됐어요!", type="positive", timeout=5000)

                except ValueError as ve:
                    thumb_status.set_text("썸네일을 만들지 못했어요")
                    ui.notify(str(ve), type="negative", timeout=8000)
                    from app.ai.image_provider import get_image_failure_guide
                    ui.notify(get_image_failure_guide(str(ve)), type="info", timeout=15000, close_button="확인")
                except Exception as exc:
                    thumb_status.set_text("썸네일을 만들지 못했어요")
                    ui.notify(f"썸네일을 만들지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative", timeout=8000)
                    from app.ai.image_provider import get_image_failure_guide
                    ui.notify(get_image_failure_guide(str(exc)), type="info", timeout=15000, close_button="확인")
                finally:
                    thumb_gen_btn.props(remove="disabled loading")
                    thumb_spinner.classes("hidden")

            def _save_thumbnail() -> None:
                from datetime import datetime
                from app.paths import THUMBNAILS_DIR
                img_bytes = _thumb_state.get("result_bytes")
                if not img_bytes:
                    ui.notify("저장할 썸네일이 아직 없어요. 먼저 썸네일을 만들어 주세요.", type="warning")
                    return
                mime = _thumb_state.get("result_mime", "image/png")
                ext = ".png" if "png" in mime else ".jpeg" if "jpeg" in mime else ".png"
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"thumbnail_{timestamp}{ext}"
                ExportManager.save_default(img_bytes, filename, dest_dir=THUMBNAILS_DIR)

            # ── Helper functions ─────────────────────────────────────────────

            def _copy_section(key: str) -> None:
                sections = _parse_planning_sections(_s2.get("content", ""))
                text = sections.get(key, "")
                if not text:
                    ui.notify("복사할 내용이 아직 없어요.", type="warning")
                    return
                escaped = json.dumps(text)
                ui.run_javascript(f'navigator.clipboard.writeText({escaped})')
                ui.notify("클립보드에 복사했어요!", type="positive", timeout=2000)

            def _render_news_card(container, text: str, label: str) -> None:
                container.clear()
                if not text.strip():
                    with container:
                        ui.label(f"*{label} 블록을 찾을 수 없어요.*").classes("dg-label-sm")
                    return

                lines = text.strip().splitlines()
                hook_text = ""
                title_text = ""
                body_lines: list[str] = []
                cta_lines: list[str] = []
                location_line = ""

                for line in lines:
                    stripped = line.strip()
                    if not stripped or stripped.startswith("\u2501") or stripped.startswith("\u3010"):
                        continue
                    if re.match(r"^\uc81c\ubaa9:\s*", stripped):
                        title_text = re.sub(r"^\uc81c\ubaa9:\s*", "", stripped)
                    elif "\ucfe0\ud3f0" in stripped and not title_text and not body_lines:
                        hook_text = stripped
                    elif re.search(r"\uc5d0\uc11c \ub9cc\ub098\uc694|\uc5d0\uc11c \uae30\ub2e4\ub9b4\uac8c\uc694|\ucc3e\uc544\uc624\uc2dc\ub294", stripped):
                        location_line = stripped
                    elif re.search(r"\ucc44\ud305|\ubb38\uc758|\ud655\uc778\ud574\s*\ubcf4\uc138\uc694|\uc0c1\ub2f4|\uc608\uc57d", stripped) and len(stripped) < 40:
                        cta_lines.append(stripped)
                    else:
                        body_lines.append(stripped)

                with container:
                    if hook_text:
                        with ui.element("div").style(
                            "background: var(--dg-primary-light); color: var(--dg-primary); "
                            "padding: 8px 16px; border-radius: 8px; font-weight: 600; font-size: 14px; text-align: center;"
                        ):
                            ui.label(hook_text)
                    if title_text:
                        with ui.element("div").style(
                            "background: var(--dg-card); border: 2px solid var(--dg-primary); "
                            "border-radius: 12px; padding: 16px 20px;"
                        ):
                            ui.label(title_text).style(
                                "font-size: 18px; font-weight: 700; color: var(--dg-text-primary); line-height: 1.4;"
                            )
                    if body_lines:
                        body_html = "<br>".join(
                            f"<p style='margin:0 0 8px 0;line-height:1.7;'>{ln}</p>" if ln else "<br>"
                            for ln in body_lines
                        )
                        with ui.element("div").style(
                            "background: var(--dg-surface); border-radius: 12px; padding: 20px; "
                            "font-size: 14px; color: var(--dg-text-secondary); line-height: 1.7;"
                        ):
                            ui.html(body_html)
                    if cta_lines:
                        with ui.row().classes("gap-2 flex-wrap"):
                            for cta in cta_lines[:3]:
                                with ui.element("div").style(
                                    "background: var(--dg-primary); color: white; "
                                    "padding: 8px 20px; border-radius: 20px; "
                                    "font-size: 13px; font-weight: 600;"
                                ):
                                    ui.label(cta)
                    if location_line:
                        with ui.row().classes("items-center gap-1"):
                            ui.icon("place", size="16px").style("color: var(--dg-text-tertiary)")
                            ui.label(location_line).classes("dg-label-sm")

            def _render_mobile_frame(container, text: str, label: str) -> None:
                container.clear()
                with container:
                    ui.label(label).style(
                        "font-size: 13px; font-weight: 600; color: var(--dg-text-secondary);"
                    )
                    with ui.element("div").style(
                        "width: 320px; min-height: 500px; max-height: 640px; overflow-y: auto; "
                        "border: 3px solid var(--dg-text-primary); border-radius: 32px; "
                        "padding: 40px 16px 24px; background: white; "
                        "box-shadow: 0 8px 32px rgba(0,0,0,0.12); position: relative;"
                    ):
                        ui.element("div").style(
                            "position: absolute; top: 10px; left: 50%; transform: translateX(-50%); "
                            "width: 80px; height: 6px; background: var(--dg-border); border-radius: 3px;"
                        )
                        clean = text.strip()
                        if not clean:
                            ui.label("아직 소식글이 없어요. 콘텐츠를 생성하면 여기에 표시돼요.").classes("dg-label-sm")
                            return
                        lines = clean.splitlines()
                        for line in lines:
                            stripped = line.strip()
                            if not stripped:
                                ui.element("div").style("height: 8px;")
                            elif stripped.startswith("\u2501") or stripped.startswith("\u3010"):
                                continue
                            elif re.match(r"^\uc81c\ubaa9:\s*", stripped):
                                title = re.sub(r"^\uc81c\ubaa9:\s*", "", stripped)
                                ui.label(title).style(
                                    "font-size: 16px; font-weight: 700; color: #212124; "
                                    "margin-bottom: 8px; line-height: 1.3;"
                                )
                            elif "\ucfe0\ud3f0" in stripped and len(stripped) < 50:
                                ui.label(stripped).style(
                                    "font-size: 13px; font-weight: 600; color: var(--dg-primary); "
                                    "margin-bottom: 4px;"
                                )
                            elif re.search(r"\ucc44\ud305|\ubb38\uc758|\ud655\uc778\ud574\s*\ubcf4\uc138\uc694", stripped) and len(stripped) < 40:
                                with ui.element("div").style(
                                    "background: var(--dg-primary); color: white; "
                                    "padding: 6px 14px; border-radius: 16px; display: inline-block; "
                                    "font-size: 12px; font-weight: 600; margin: 4px 0;"
                                ):
                                    ui.label(stripped)
                            else:
                                ui.label(stripped).style(
                                    "font-size: 13px; color: #4E5968; line-height: 1.6; margin-bottom: 2px;"
                                )

            def _update_result_display(content: str) -> None:
                sections = _parse_planning_sections(content)
                result_md_all.set_content(content)
                _render_news_card(result_card_v1, sections["version_1"], "의심해소형")
                _render_news_card(result_card_v2, sections["version_2"], "가성비형")
                result_md_copies.set_content(sections.get("ad_copies", "") or "*카피 섹션을 찾을 수 없어요.*")
                result_md_campaign.set_content(sections.get("campaign_groups", "") or "*캠페인 그룹 섹션을 찾을 수 없어요.*")
                result_md_thumbnail.set_content(sections.get("thumbnail_guide", "") or "*썸네일 가이드를 찾을 수 없어요.*")
                _coupon_naming = ""
                if sections.get("coupon_spec"):
                    _coupon_naming += "### 쿠폰 스펙\n" + sections["coupon_spec"]
                if sections.get("naming_convention"):
                    _coupon_naming += "\n### 캠페인 네이밍\n" + sections["naming_convention"]
                result_md_coupon.set_content(_coupon_naming or "*쿠폰/네이밍 섹션을 찾을 수 없어요.*")
                _render_mobile_frame(mobile_frame_v1, sections["version_1"], "소식글 1 | 의심해소형")
                _render_mobile_frame(mobile_frame_v2, sections["version_2"], "소식글 2 | 가성비형")
                # Update thumbnail generation tab guide
                result_md_thumbnail_gen.set_content(
                    sections.get("thumbnail_guide", "") or "*썸네일 가이드를 찾을 수 없어요.*"
                )
                result_card.classes(remove="hidden")
                feedback_card_s2.classes(remove="hidden")

            # ── Generate content ─────────────────────────────────────────────

            async def _generate_content() -> None:
                pid = nicegui_app.storage.user.get("current_project_id")
                if not pid:
                    ui.notify("프로젝트를 먼저 선택해 주세요.", type="warning")
                    return
                project = get_project(pid)
                if not project:
                    ui.notify("프로젝트를 찾을 수 없어요. 프로젝트 페이지에서 다시 선택해 주세요.", type="negative")
                    return

                engine = engine_radio.value
                extra = extra_input.value
                cat = category_sel.value or "default"
                strat = strategy_sel.value or "A"

                _s2["cancelled"] = False
                spinner.classes(remove="hidden")
                gen_btn.props("disabled")
                cancel_btn.classes(remove="hidden")

                # Validate API keys
                # Claude는 기본이 CLI(구독). API 백엔드일 때만 키 필요.
                if (engine == "claude" and os.getenv("CLAUDE_BACKEND", "cli").strip().lower() == "api"
                        and not os.getenv("ANTHROPIC_API_KEY", "")):
                    ui.notify("Claude를 API 모드로 쓰려면 .env에 ANTHROPIC_API_KEY가 필요해요. (CLI 모드는 키 없이 동작)", type="negative")
                    spinner.classes("hidden")
                    gen_btn.props(remove="disabled")
                    cancel_btn.classes("hidden")
                    return
                if (engine in ("gpt", "coordinate") and os.getenv("OPENAI_BACKEND", "cli").strip().lower() == "api"
                        and not os.getenv("OPENAI_API_KEY", "")):
                    ui.notify("GPT를 API 모드로 쓰려면 .env에 OPENAI_API_KEY가 필요해요. (CLI 모드는 codex 로그인으로 동작)", type="negative")
                    spinner.classes("hidden")
                    gen_btn.props(remove="disabled")
                    cancel_btn.classes("hidden")
                    return

                try:
                    _set_step("1/3 콘텐츠 작성을 준비하고 있어요...")
                    strategy_ctx = _wizard_state.get("step1_content", "")
                    guide, prompt = build_planning_prompt(
                        project, extra, category=cat, strategy=strat,
                        engine=engine if engine not in ("both", "coordinate") else "",
                        strategy_context=strategy_ctx,
                    )
                    _custom = get_setting("custom_system_prompt")
                    if _custom:
                        guide = _custom
                    loop = asyncio.get_running_loop()

                    if _s2["cancelled"]:
                        ui.notify("생성을 중단했어요.", type="warning")
                        return

                    if engine == "coordinate":
                        from app.ai.coordination import synthesize
                        _set_step("2/3 Claude와 GPT가 각자 초안을 쓰고 있어요...")
                        claude_guide, _ = build_planning_prompt(
                            project, extra, category=cat, strategy=strat, engine="claude",
                            strategy_context=strategy_ctx,
                        )
                        gpt_guide, _ = build_planning_prompt(
                            project, extra, category=cat, strategy=strat, engine="gpt",
                            strategy_context=strategy_ctx,
                        )
                        if _custom:
                            claude_guide = _custom
                            gpt_guide = _custom
                        claude_p = get_provider("claude")
                        gpt_p = get_provider("gpt")
                        c_text, g_text = await asyncio.gather(
                            loop.run_in_executor(None, lambda: claude_p.generate_text(prompt, system_prompt=claude_guide)),
                            loop.run_in_executor(None, lambda: gpt_p.generate_text(prompt, system_prompt=gpt_guide)),
                        )
                        if _s2["cancelled"]:
                            ui.notify("생성을 중단했어요.", type="warning")
                            return
                        _set_step("3/3 Claude가 두 초안을 종합하고 있어요...")
                        content = await loop.run_in_executor(
                            None, lambda: synthesize(c_text, g_text, "기획 콘텐츠 생성"),
                        )
                    else:
                        engine_name = "GPT" if engine == "gpt" else "Claude"
                        _set_step(f"2/3 {engine_name}가 콘텐츠를 작성하고 있어요...")
                        provider = get_provider(engine)
                        content = await loop.run_in_executor(None, lambda: provider.generate_text(prompt, system_prompt=guide))
                        if _s2["cancelled"]:
                            ui.notify("생성을 중단했어요.", type="warning")
                            return

                    _set_step("3/3 결과를 저장하고 있어요...")
                    _s2["content"] = content
                    _s2["engine"] = engine
                    _wizard_state["step2_content"] = content
                    _wizard_state["step2_engine"] = engine
                    save_generated_content(pid, engine, content)

                    _update_result_display(content)
                    ui.notify("기획 콘텐츠가 완성됐어요!", type="positive")

                except Exception as exc:
                    _log.exception("기획 콘텐츠 생성 실패: %s", exc)
                    ui.notify(f"콘텐츠를 만들지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative", timeout=12000, close_button="확인")
                finally:
                    spinner.classes("hidden")
                    cancel_btn.classes("hidden")
                    step_label.classes("hidden")
                    gen_btn.props(remove="disabled")

            # ── Feedback & regenerate (Step 2) ────────────────────────────────

            feedback_card_s2 = ui.card().classes("dg-card-flat w-full hidden")
            with feedback_card_s2:
                ui.label("수정 요청").style("font-weight: 600; font-size: 14px; color: var(--dg-text-primary)")
                feedback_s2 = ui.textarea(
                    placeholder="예: 소식글 톤을 더 캐주얼하게, 쿠폰 혜택을 강조해주세요 등"
                ).classes("w-full dg-input").props("rows=2 outlined")

                with ui.row().classes("gap-3 items-center"):
                    regen_s2_btn = ui.button(
                        "재생성", icon="refresh",
                        on_click=lambda: _regenerate_content(feedback_s2.value),
                    ).classes("dg-btn-ghost")
                    regen_s2_spinner = ui.spinner(size="24px").classes("hidden")
                    regen_s2_status = ui.label("").classes("dg-progress-text hidden")

                _wizard_state["_regen_s2_spinner"] = regen_s2_spinner
                _wizard_state["_regen_s2_status"] = regen_s2_status
                _wizard_state["_regen_s2_btn"] = regen_s2_btn

            async def _regenerate_content(feedback: str = "") -> None:
                pid = nicegui_app.storage.user.get("current_project_id")
                if not pid:
                    return
                project = get_project(pid)
                if not project:
                    return

                engine = engine_radio.value
                # 재생성은 단일 모델로 처리(조율은 비용 큼) — 종합 엔진(claude)로 collapse
                if engine == "coordinate":
                    engine = "claude"

                regen_spinner = _wizard_state.get("_regen_s2_spinner")
                regen_status = _wizard_state.get("_regen_s2_status")
                regen_btn = _wizard_state.get("_regen_s2_btn")

                if regen_btn:
                    regen_btn.props("disabled loading")
                if regen_spinner:
                    regen_spinner.classes(remove="hidden")
                if regen_status:
                    regen_status.classes(remove="hidden")
                    regen_status.set_text("다시 만들고 있어요...")

                try:
                    cat = category_sel.value or "default"
                    strat = strategy_sel.value or "A"
                    strategy_ctx = _wizard_state.get("step1_content", "")
                    guide, prompt = build_planning_prompt(
                        project, extra_input.value, category=cat, strategy=strat,
                        engine=engine, strategy_context=strategy_ctx,
                    )
                    _custom = get_setting("custom_system_prompt")
                    if _custom:
                        guide = _custom
                    if feedback.strip():
                        prompt += f"\n\n[수정 요청]\n{feedback.strip()}"
                    if _wizard_state["step2_content"]:
                        prompt += f"\n\n[이전 결과 (참고하되 수정 요청 반영)]\n{_wizard_state['step2_content']}"

                    loop = asyncio.get_running_loop()
                    provider = get_provider(engine)
                    content = await loop.run_in_executor(
                        None, lambda: provider.generate_text(prompt, system_prompt=guide),
                    )

                    _s2["content"] = content
                    _s2["engine"] = engine
                    _wizard_state["step2_content"] = content
                    _wizard_state["step2_engine"] = engine
                    save_generated_content(pid, engine, content)

                    _update_result_display(content)
                    ui.notify("콘텐츠를 다시 만들었어요!", type="positive")

                except Exception as exc:
                    _log.exception("콘텐츠 재생성 실패: %s", exc)
                    ui.notify(f"다시 만들지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative", timeout=8000)
                finally:
                    if regen_btn:
                        regen_btn.props(remove="disabled loading")
                    if regen_spinner:
                        regen_spinner.classes("hidden")

            # ── Export handlers ───────────────────────────────────────────────

            def _build_docx_bytes(pid_snapshot: str | None = None) -> tuple[bytes, str]:
                import tempfile
                content = _s2.get("content", "")
                if not content:
                    raise ValueError("콘텐츠를 먼저 만들어 주세요.")
                project = get_project(pid_snapshot) if pid_snapshot else None
                if not project:
                    raise ValueError("프로젝트를 선택해 주세요.")
                name = project.get("name", "unknown")
                project_meta = {
                    k: project.get(k, "")
                    for k in ("name", "period", "goal", "industry", "region", "budget")
                }
                with tempfile.TemporaryDirectory() as tmp_dir:
                    tmp_path = Path(tmp_dir) / f"기획서_{name}.docx"
                    build_planning_docx(project_meta, content, tmp_path)
                    return tmp_path.read_bytes(), f"기획서_{name}.docx"

            async def _export_default() -> None:
                download_status.classes(remove="hidden")
                download_status.set_text("DOCX 파일을 준비하고 있어요...")
                try:
                    pid_snap = nicegui_app.storage.user.get("current_project_id")
                    loop = asyncio.get_running_loop()
                    docx_bytes, fname = await loop.run_in_executor(None, lambda: _build_docx_bytes(pid_snap))
                    ExportManager.save_default(docx_bytes, fname)
                    download_status.set_text(f"{fname} 저장 완료")
                    ui.notify(f"{fname}", type="positive", timeout=8000, close_button="확인")
                except ValueError as ve:
                    ui.notify(str(ve), type="warning")
                except Exception as exc:
                    download_status.set_text("내보내지 못했어요")
                    ui.notify(f"파일을 내보내지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative")

            async def _export_default_pdf() -> None:
                download_status.classes(remove="hidden")
                download_status.set_text("PDF 파일을 준비하고 있어요...")
                try:
                    import tempfile
                    pid_snap = nicegui_app.storage.user.get("current_project_id")
                    loop = asyncio.get_running_loop()
                    docx_bytes, fname = await loop.run_in_executor(None, lambda: _build_docx_bytes(pid_snap))
                    pdf_fname = fname.replace(".docx", ".pdf")

                    def _convert_to_pdf():
                        from docx2pdf import convert
                        import pythoncom
                        pythoncom.CoInitialize()
                        try:
                            with tempfile.TemporaryDirectory() as tmpdir:
                                docx_path = Path(tmpdir) / fname
                                docx_path.write_bytes(docx_bytes)
                                pdf_path = Path(tmpdir) / pdf_fname
                                convert(str(docx_path), str(pdf_path))
                                return pdf_path.read_bytes()
                        finally:
                            pythoncom.CoUninitialize()

                    try:
                        pdf_bytes = await loop.run_in_executor(None, _convert_to_pdf)
                        ExportManager.save_default(pdf_bytes, pdf_fname)
                        download_status.set_text(f"{pdf_fname} 저장 완료")
                        ui.notify(f"{pdf_fname}", type="positive", timeout=8000, close_button="확인")
                    except ImportError:
                        ui.notify("PDF로 저장하려면 docx2pdf 설치가 필요해요. 터미널에서 pip install docx2pdf를 실행해 주세요.", type="warning")
                        download_status.set_text("PDF 변환 불가 - docx2pdf 없음")
                    except Exception as pdf_exc:
                        ui.notify(f"PDF로 바꾸지 못해서 DOCX로 저장했어요. ({pdf_exc})", type="warning")
                        ExportManager.save_default(docx_bytes, fname)
                        download_status.set_text(f"{fname} 저장 완료 (DOCX 대체)")
                except ValueError as ve:
                    ui.notify(str(ve), type="warning")
                except Exception as exc:
                    download_status.set_text("내보내지 못했어요")
                    ui.notify(f"파일을 내보내지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative")

            # ── Load saved content ───────────────────────────────────────────

            pid0 = nicegui_app.storage.user.get("current_project_id")
            if pid0:
                saved = get_latest_content(pid0)
                if saved:
                    _s2["content"] = saved["content"]
                    _s2["engine"] = saved.get("engine", "claude")
                    _wizard_state["step2_content"] = saved["content"]
                    _update_result_display(saved["content"])

            # Navigation buttons
            _render_export_bar(
                2,
                docx_handler=lambda: _export_default(),
                pdf_handler=lambda: _export_default_pdf(),
            )

    def _advance_to_step3() -> None:
        if not _wizard_state["step2_content"]:
            ui.notify("콘텐츠를 먼저 만들어 주세요.", type="warning")
            return
        _wizard_state["current_step"] = 3
        _render_step_indicator()
        _render_current_step()

    # ── Step 3: Ad Settings Guide ────────────────────────────────────────────

    def _render_step3() -> None:
        content_container.clear()
        with content_container:
            section_header("ads_click", "광고 세팅 가이드", "전략과 콘텐츠를 바탕으로 당근 전문가모드 광고 세팅 가이드를 만들어요.")

            # Previous steps context (collapsible)
            if _wizard_state["step1_content"]:
                with ui.expansion(
                    "전략 분석 요약 (Step 1 결과)", icon="analytics",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step1_content"]).classes("w-full dg-prose")

            if _wizard_state["step2_content"]:
                with ui.expansion(
                    "콘텐츠 요약 (Step 2 결과)", icon="edit_note",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step2_content"][:2000] + ("\n\n..." if len(_wizard_state["step2_content"]) > 2000 else "")).classes("w-full dg-prose")

            # Show existing result if loaded
            if _wizard_state["step3_content"]:
                _render_ad_settings_result()
                return

            _render_budget_planner_card()

            with ui.card().classes("dg-card w-full"):
                with ui.column().classes("w-full gap-3 items-center"):
                    ui.icon("ads_click", size="48px").style("color: var(--dg-border)")
                    ui.label(
                        "전략 분석과 콘텐츠 결과를 바탕으로 광고 세팅 가이드를 만들어 드려요."
                    ).classes("dg-text-sm").style("text-align: center")
                    ui.label(
                        "캠페인 구조, 타겟팅, 예산 배분, 소재 배치, 성과 측정 계획이 담겨요."
                    ).classes("dg-label-sm").style("text-align: center")

                    with ui.row().classes("gap-3 items-center"):
                        gen_s3_btn = ui.button(
                            "광고 세팅 가이드 생성", icon="auto_awesome",
                            on_click=lambda: _generate_ad_settings(),
                        ).classes("dg-btn-primary")
                        s3_spinner = ui.spinner(size="28px").classes("hidden")
                        s3_status = ui.label("").classes("dg-progress-text hidden")

                    _wizard_state["_s3_btn"] = gen_s3_btn
                    _wizard_state["_s3_spinner"] = s3_spinner
                    _wizard_state["_s3_status"] = s3_status

            # Back button
            with ui.row().classes("w-full justify-start mt-2"):
                ui.button(
                    "이전 단계로", icon="arrow_back",
                    on_click=lambda: _go_to_step(2),
                ).classes("dg-btn-secondary")

    def _render_budget_planner_card() -> None:
        """예산 기반 캠페인 설계 카드 (룰 엔진 — AI 호출 없음).

        일예산을 넣으면 자동+수동 페어 가능 여부와 캠페인 세팅표를 즉시 계산한다.
        계산 결과는 wizard state에 저장되어 Step 3 AI 프롬프트에 그대로 주입된다.
        """
        from app.engine.budget_planner import (
            BROAD_AGE_BAND,
            DEFAULT_AGE_BANDS,
            SETTING_TABLE_COLUMNS,
            SIMPLE_AGE_BANDS,
            plan_table_rows,
            plan_to_prompt_context,
            recommend_structure,
        )

        pid = nicegui_app.storage.user.get("current_project_id")
        proj = get_project(pid) if pid else None
        raw_region = ((proj or {}).get("region") or "").strip()
        default_region = raw_region.split()[-1] if raw_region else "우리동네"

        with ui.card().classes("dg-card w-full"):
            section_header(
                "calculate", "예산 기반 캠페인 설계",
                "일예산을 넣으면 자동+수동 페어가 가능한지 계산해서 세팅표를 만들어 드려요. AI는 이 설계를 그대로 따라요.",
            )
            with ui.row().classes("w-full gap-4 items-end flex-wrap"):
                with ui.column().classes("gap-1"):
                    ui.label("일예산 (원)").classes("dg-label-sm")
                    plan_budget_in = ui.number(
                        value=30000, min=0, step=5000, format="%d",
                    ).props("outlined dense").classes("w-32 dg-input")
                with ui.column().classes("gap-1"):
                    ui.label("지역 (캠페인명용)").classes("dg-label-sm")
                    plan_region_in = ui.input(value=default_region).props(
                        "outlined dense"
                    ).classes("w-32 dg-input")
                with ui.column().classes("gap-1"):
                    ui.label("성별").classes("dg-label-sm")
                    plan_gender_sel = ui.select(
                        ["여성", "남성", "전체"], value="여성",
                    ).props("outlined dense").classes("w-24 dg-select")
                with ui.column().classes("gap-1"):
                    ui.label("핵심 연령대").classes("dg-label-sm")
                    plan_age_sel = ui.select(
                        [BROAD_AGE_BAND] + SIMPLE_AGE_BANDS + DEFAULT_AGE_BANDS,
                        value="40대",
                    ).props("outlined dense").classes("w-28 dg-select")
                with ui.column().classes("gap-1"):
                    ui.label("소구 키워드 (캠페인명용)").classes("dg-label-sm")
                    plan_appeal_in = ui.input(value="핵심소구").props(
                        "outlined dense"
                    ).classes("w-32 dg-input")
                with ui.column().classes("gap-1"):
                    ui.label("목표 문의당 비용 (선택)").classes("dg-label-sm")
                    plan_cpa_in = ui.number(
                        value=None, min=0, step=1000, format="%d",
                    ).props("outlined dense").classes("w-32 dg-input")
                plan_validated_sw = ui.switch("검증된 소재/타겟 있음")
                ui.button(
                    "설계 계산", icon="calculate",
                    on_click=lambda: _calc_budget_plan(),
                ).classes("dg-btn-primary dg-btn-sm")

            plan_result_box = ui.column().classes("w-full gap-2 mt-2")

        def _calc_budget_plan() -> None:
            try:
                budget = int(plan_budget_in.value or 0)
                target_cpa = int(plan_cpa_in.value) if plan_cpa_in.value else None
                plan = recommend_structure(
                    budget,
                    region=(plan_region_in.value or "우리동네").strip() or "우리동네",
                    gender=plan_gender_sel.value or "여성",
                    age_band=plan_age_sel.value or "40대",
                    appeal=(plan_appeal_in.value or "핵심소구").strip() or "핵심소구",
                    has_validated_creative=bool(plan_validated_sw.value),
                    target_cpa=target_cpa,
                )
            except Exception as exc:
                ui.notify(f"설계를 계산하지 못했어요. 입력값을 확인해 주세요. ({exc})", type="negative")
                return

            _wizard_state["budget_plan_text"] = plan_to_prompt_context(plan)

            plan_result_box.clear()
            with plan_result_box:
                banner_cls = (
                    "dg-banner-warning"
                    if plan.tier in ("below_minimum", "single")
                    else "dg-banner-info"
                )
                with ui.element("div").classes(f"dg-banner {banner_cls} w-full"):
                    ui.icon("insights", size="18px")
                    with ui.column().classes("gap-1"):
                        ui.label(plan.mode_label).style("font-weight: 700; font-size: 13px")
                        ui.label(plan.feasibility_note).style("font-size: 12px")
                for w in plan.warnings:
                    ui.label(w).classes("dg-label-sm").style("color: var(--dg-warning)")
                if plan.campaigns:
                    ui.table(
                        columns=SETTING_TABLE_COLUMNS,
                        rows=plan_table_rows(plan),
                        row_key="name",
                    ).classes("w-full dg-table").props("dense flat bordered")
                if plan.next_steps:
                    ui.label("다음 단계").style(
                        "font-size: 13px; font-weight: 600; margin-top: 4px"
                    )
                    for s in plan.next_steps:
                        ui.label(f"- {s}").classes("dg-label-sm")
                ui.label(
                    "이 설계는 '광고 세팅 가이드 생성'을 누르면 AI에게 그대로 전달돼요."
                ).classes("dg-label-sm").style("color: var(--dg-text-tertiary)")
            ui.notify("캠페인 설계를 계산했어요. 세팅표를 확인해 보세요.", type="positive")

    def _render_ad_settings_result() -> None:
        """Render ad settings results with edit/regenerate controls."""
        content_container.clear()
        sections = _wizard_state["step3_sections"]

        with content_container:
            section_header("ads_click", "광고 세팅 가이드 결과")

            _render_export_bar(3)

            # Previous steps context (collapsible)
            if _wizard_state["step1_content"]:
                with ui.expansion(
                    "전략 분석 요약 (Step 1 결과)", icon="analytics",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step1_content"]).classes("w-full dg-prose")

            if _wizard_state["step2_content"]:
                with ui.expansion(
                    "콘텐츠 요약 (Step 2 결과)", icon="edit_note",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step2_content"][:2000] + ("\n\n..." if len(_wizard_state["step2_content"]) > 2000 else "")).classes("w-full dg-prose")

            # 결과가 있어도 예산 설계를 다시 계산해 재생성에 반영할 수 있게 노출
            with ui.expansion(
                "예산 기반 캠페인 설계 (계산해서 재생성에 반영)", icon="calculate",
            ).classes("w-full dg-expansion").props("dense"):
                _render_budget_planner_card()

            for key in _AD_SETTINGS_SECTION_KEYS:
                label = _AD_SETTINGS_SECTION_NAMES[key]
                body = sections.get(key, "(아직 내용이 없어요)")
                idx = _AD_SETTINGS_SECTION_KEYS.index(key) + 1

                with ui.expansion(
                    f"{idx}. {label}",
                    icon="article",
                    value=True,
                ).classes("w-full dg-expansion"):
                    md_widget = ui.markdown(body).classes("w-full dg-prose")
                    edit_area = ui.textarea(value=body).classes(
                        "w-full hidden dg-input"
                    ).props("rows=8 outlined")

                    with ui.row().classes("gap-2 mt-2"):
                        edit_btn = ui.button("편집", icon="edit").classes("dg-btn-ghost dg-btn-sm")
                        save_btn = ui.button("편집 완료", icon="check").classes("dg-btn-success dg-btn-sm hidden")

                        def _toggle_edit_s3(
                            _e, _md=md_widget, _ea=edit_area, _eb=edit_btn, _sb=save_btn,
                        ) -> None:
                            _md.classes("hidden", remove=False)
                            _ea.classes(remove="hidden")
                            _eb.classes("hidden", remove=False)
                            _sb.classes(remove="hidden")

                        def _save_edit_s3(
                            _e, _md=md_widget, _ea=edit_area, _eb=edit_btn, _sb=save_btn,
                            _key=key,
                        ) -> None:
                            new_text = _ea.value
                            _wizard_state["step3_sections"][_key] = new_text
                            _md.set_content(new_text)
                            _md.classes(remove="hidden")
                            _ea.classes("hidden", remove=False)
                            _sb.classes("hidden", remove=False)
                            _eb.classes(remove="hidden")
                            parts = []
                            for i, k in enumerate(_AD_SETTINGS_SECTION_KEYS):
                                n = _AD_SETTINGS_SECTION_NAMES[k]
                                b = _wizard_state["step3_sections"].get(k, "")
                                parts.append(f"## {i + 1}. {n}\n{b}")
                            _wizard_state["step3_content"] = "\n\n".join(parts)
                            pid = nicegui_app.storage.user.get("current_project_id")
                            if pid:
                                save_generated_content(pid, "edited", _wizard_state["step3_content"], content_type="ad_settings")
                            ui.notify("수정한 내용을 저장했어요.", type="positive")

                        edit_btn.on_click(_toggle_edit_s3)
                        save_btn.on_click(_save_edit_s3)

            # Feedback + regenerate
            with ui.card().classes("dg-card-flat w-full"):
                ui.label("수정 요청").style("font-weight: 600; font-size: 14px; color: var(--dg-text-primary)")
                feedback_s3 = ui.textarea(
                    placeholder="예: 예산 배분을 좀 더 보수적으로, A/B 테스트 기간을 2주로 변경 등"
                ).classes("w-full dg-input").props("rows=2 outlined")

                with ui.row().classes("gap-3 items-center"):
                    regen_s3_btn = ui.button(
                        "재생성", icon="refresh",
                        on_click=lambda: _regenerate_ad_settings(feedback_s3.value),
                    ).classes("dg-btn-ghost")
                    regen_s3_spinner = ui.spinner(size="24px").classes("hidden")
                    regen_s3_status = ui.label("").classes("dg-progress-text hidden")

                _wizard_state["_regen_s3_spinner"] = regen_s3_spinner
                _wizard_state["_regen_s3_status"] = regen_s3_status
                _wizard_state["_regen_s3_btn"] = regen_s3_btn

            # Navigation
            _render_export_bar(3)

    async def _generate_ad_settings() -> None:
        pid = nicegui_app.storage.user.get("current_project_id")
        if not pid:
            ui.notify("프로젝트를 먼저 선택해 주세요.", type="warning")
            return
        project = get_project(pid)
        if not project:
            ui.notify("프로젝트를 찾을 수 없어요. 프로젝트 페이지에서 다시 선택해 주세요.", type="negative")
            return

        engine = engine_radio.value

        # Claude는 기본이 CLI(구독). API 백엔드일 때만 키 필요.
        if (engine == "claude" and os.getenv("CLAUDE_BACKEND", "cli").strip().lower() == "api"
                and not os.getenv("ANTHROPIC_API_KEY", "")):
            ui.notify("Claude를 API 모드로 쓰려면 .env에 ANTHROPIC_API_KEY가 필요해요. (CLI 모드는 키 없이 동작)", type="negative")
            return
        # GPT는 기본이 codex CLI(구독). API 백엔드일 때만 키 필요.
        if (engine in ("gpt", "coordinate") and os.getenv("OPENAI_BACKEND", "cli").strip().lower() == "api"
                and not os.getenv("OPENAI_API_KEY", "")):
            ui.notify("GPT를 API 모드로 쓰려면 .env에 OPENAI_API_KEY가 필요해요. (CLI 모드는 codex 로그인으로 동작)", type="negative")
            return

        btn = _wizard_state.get("_s3_btn")
        spinner = _wizard_state.get("_s3_spinner")
        status = _wizard_state.get("_s3_status")

        if btn:
            btn.props("disabled loading")
        if spinner:
            spinner.classes(remove="hidden")
        if status:
            status.classes(remove="hidden")
            status.set_text("광고 세팅 가이드를 만들고 있어요...")

        try:
            guide, prompt = build_ad_settings_prompt(
                project,
                strategy_context=_wizard_state.get("step1_content", ""),
                content_context=_wizard_state.get("step2_content", ""),
                budget_plan_context=_wizard_state.get("budget_plan_text", ""),
            )
            loop = asyncio.get_running_loop()
            if engine == "coordinate":
                from app.ai.coordination import coordinate_generate
                content = await coordinate_generate(
                    loop, prompt, guide, "광고 세팅 가이드",
                    on_drafts=(lambda: status.set_text("Claude와 GPT가 각자 작성하고 있어요...")) if status else None,
                    on_synth=(lambda: status.set_text("Claude가 두 초안을 종합하고 있어요...")) if status else None,
                )
            else:
                provider = get_provider(engine)
                content = await loop.run_in_executor(
                    None, lambda: provider.generate_text(prompt, system_prompt=guide),
                )

            # 누락/부실 섹션이 있으면 1회 보정 (best-effort, 실패 시 원본 유지).
            if status:
                status.set_text("빠진 부분이 없는지 확인하고 있어요...")
            content = await loop.run_in_executor(
                None, lambda: repair_output(content, get_schema("ad_settings"), engine=engine),
            )

            _wizard_state["step3_content"] = content
            _wizard_state["step3_sections"] = parse_ad_settings_sections(content)
            save_generated_content(pid, engine, content, content_type="ad_settings")

            _render_ad_settings_result()
            ui.notify("광고 세팅 가이드가 완성됐어요!", type="positive")

        except Exception as exc:
            _log.exception("광고 세팅 가이드 생성 실패: %s", exc)
            ui.notify(f"광고 세팅 가이드를 만들지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative", timeout=8000)
            if status:
                status.set_text(f"오류: {exc}")
        finally:
            if btn:
                btn.props(remove="disabled loading")
            if spinner:
                spinner.classes("hidden")

    async def _regenerate_ad_settings(feedback: str = "") -> None:
        pid = nicegui_app.storage.user.get("current_project_id")
        if not pid:
            return
        project = get_project(pid)
        if not project:
            return

        engine = engine_radio.value if engine_radio.value not in ("both", "coordinate") else "claude"

        regen_spinner = _wizard_state.get("_regen_s3_spinner")
        regen_status = _wizard_state.get("_regen_s3_status")
        regen_btn = _wizard_state.get("_regen_s3_btn")

        if regen_btn:
            regen_btn.props("disabled loading")
        if regen_spinner:
            regen_spinner.classes(remove="hidden")
        if regen_status:
            regen_status.classes(remove="hidden")
            regen_status.set_text("다시 만들고 있어요...")

        try:
            guide, prompt = build_ad_settings_prompt(
                project,
                strategy_context=_wizard_state.get("step1_content", ""),
                content_context=_wizard_state.get("step2_content", ""),
                budget_plan_context=_wizard_state.get("budget_plan_text", ""),
            )
            if feedback.strip():
                prompt += f"\n\n[수정 요청]\n{feedback.strip()}"
            if _wizard_state["step3_content"]:
                prompt += f"\n\n[이전 결과 (참고하되 수정 요청 반영)]\n{_wizard_state['step3_content']}"

            loop = asyncio.get_running_loop()
            provider = get_provider(engine)
            content = await loop.run_in_executor(
                None, lambda: provider.generate_text(prompt, system_prompt=guide),
            )

            _wizard_state["step3_content"] = content
            _wizard_state["step3_sections"] = parse_ad_settings_sections(content)
            save_generated_content(pid, engine, content, content_type="ad_settings")

            _render_ad_settings_result()
            ui.notify("광고 세팅 가이드를 다시 만들었어요!", type="positive")

        except Exception as exc:
            _log.exception("광고 세팅 가이드 재생성 실패: %s", exc)
            ui.notify(f"다시 만들지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative", timeout=8000)
        finally:
            if regen_btn:
                regen_btn.props(remove="disabled loading")
            if regen_spinner:
                regen_spinner.classes("hidden")

    def _advance_to_step4() -> None:
        if not _wizard_state["step3_content"]:
            ui.notify("광고 세팅 가이드를 먼저 만들어 주세요.", type="warning")
            return
        _wizard_state["current_step"] = 4
        _render_step_indicator()
        _render_current_step()

    # ── Step 4: Operational Proposal ─────────────────────────────────────────

    def _render_step4() -> None:
        content_container.clear()
        with content_container:
            section_header("description", "운영 제안서", "전략, 콘텐츠, 광고 세팅을 종합한 통합 운영 제안서를 만들어요.")

            # Previous steps context (collapsible)
            if _wizard_state["step1_content"]:
                with ui.expansion(
                    "전략 분석 요약 (Step 1 결과)", icon="analytics",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step1_content"]).classes("w-full dg-prose")

            if _wizard_state["step2_content"]:
                with ui.expansion(
                    "콘텐츠 요약 (Step 2 결과)", icon="edit_note",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step2_content"][:2000] + ("\n\n..." if len(_wizard_state["step2_content"]) > 2000 else "")).classes("w-full dg-prose")

            if _wizard_state["step3_content"]:
                with ui.expansion(
                    "광고 세팅 요약 (Step 3 결과)", icon="ads_click",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step3_content"]).classes("w-full dg-prose")

            # Show existing result if loaded
            if _wizard_state["step4_content"]:
                _render_proposal_result()
                return

            with ui.card().classes("dg-card w-full"):
                with ui.column().classes("w-full gap-3 items-center"):
                    ui.icon("description", size="48px").style("color: var(--dg-border)")
                    ui.label(
                        "모든 분석 결과를 종합해 통합 운영 제안서를 만들어 드려요."
                    ).classes("dg-text-sm").style("text-align: center")
                    ui.label(
                        "요약, 타겟, 콘텐츠, 광고 집행, 예산/KPI, 운영 일정, 성과 판단 기준이 담겨요."
                    ).classes("dg-label-sm").style("text-align: center")

                    with ui.row().classes("gap-3 items-center"):
                        gen_s4_btn = ui.button(
                            "운영 제안서 생성", icon="auto_awesome",
                            on_click=lambda: _generate_proposal(),
                        ).classes("dg-btn-primary")
                        s4_spinner = ui.spinner(size="28px").classes("hidden")
                        s4_status = ui.label("").classes("dg-progress-text hidden")

                    _wizard_state["_s4_btn"] = gen_s4_btn
                    _wizard_state["_s4_spinner"] = s4_spinner
                    _wizard_state["_s4_status"] = s4_status

            # Back button
            with ui.row().classes("w-full justify-start mt-2"):
                ui.button(
                    "이전 단계로", icon="arrow_back",
                    on_click=lambda: _go_to_step(3),
                ).classes("dg-btn-secondary")

    def _render_proposal_result() -> None:
        """Render proposal results with edit/regenerate/export controls."""
        content_container.clear()
        sections = _wizard_state["step4_sections"]

        with content_container:
            section_header("description", "운영 제안서 결과")

            _render_export_bar(
                4,
                docx_handler=lambda: _export_proposal_default(),
                pdf_handler=lambda: _export_proposal_default_pdf(),
            )

            # Previous steps context (collapsible)
            if _wizard_state["step1_content"]:
                with ui.expansion(
                    "전략 분석 요약 (Step 1 결과)", icon="analytics",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step1_content"]).classes("w-full dg-prose")

            if _wizard_state["step2_content"]:
                with ui.expansion(
                    "콘텐츠 요약 (Step 2 결과)", icon="edit_note",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step2_content"][:2000] + ("\n\n..." if len(_wizard_state["step2_content"]) > 2000 else "")).classes("w-full dg-prose")

            if _wizard_state["step3_content"]:
                with ui.expansion(
                    "광고 세팅 요약 (Step 3 결과)", icon="ads_click",
                ).classes("w-full dg-expansion").props("dense"):
                    ui.markdown(_wizard_state["step3_content"]).classes("w-full dg-prose")

            for key in _WIZARD_PROPOSAL_SECTION_KEYS:
                label = _WIZARD_PROPOSAL_SECTION_NAMES[key]
                body = sections.get(key, "(아직 내용이 없어요)")
                idx = _WIZARD_PROPOSAL_SECTION_KEYS.index(key) + 1

                with ui.expansion(
                    f"{idx}. {label}",
                    icon="article",
                    value=True,
                ).classes("w-full dg-expansion"):
                    md_widget = ui.markdown(body).classes("w-full dg-prose")
                    edit_area = ui.textarea(value=body).classes(
                        "w-full hidden dg-input"
                    ).props("rows=8 outlined")

                    with ui.row().classes("gap-2 mt-2"):
                        edit_btn = ui.button("편집", icon="edit").classes("dg-btn-ghost dg-btn-sm")
                        save_btn = ui.button("편집 완료", icon="check").classes("dg-btn-success dg-btn-sm hidden")

                        def _toggle_edit_s4(
                            _e, _md=md_widget, _ea=edit_area, _eb=edit_btn, _sb=save_btn,
                        ) -> None:
                            _md.classes("hidden", remove=False)
                            _ea.classes(remove="hidden")
                            _eb.classes("hidden", remove=False)
                            _sb.classes(remove="hidden")

                        def _save_edit_s4(
                            _e, _md=md_widget, _ea=edit_area, _eb=edit_btn, _sb=save_btn,
                            _key=key,
                        ) -> None:
                            new_text = _ea.value
                            _wizard_state["step4_sections"][_key] = new_text
                            _md.set_content(new_text)
                            _md.classes(remove="hidden")
                            _ea.classes("hidden", remove=False)
                            _sb.classes("hidden", remove=False)
                            _eb.classes(remove="hidden")
                            parts = []
                            for i, k in enumerate(_WIZARD_PROPOSAL_SECTION_KEYS):
                                n = _WIZARD_PROPOSAL_SECTION_NAMES[k]
                                b = _wizard_state["step4_sections"].get(k, "")
                                parts.append(f"## {i + 1}. {n}\n{b}")
                            _wizard_state["step4_content"] = "\n\n".join(parts)
                            pid = nicegui_app.storage.user.get("current_project_id")
                            if pid:
                                save_generated_content(pid, "edited", _wizard_state["step4_content"], content_type="wizard_proposal")
                            ui.notify("수정한 내용을 저장했어요.", type="positive")

                        edit_btn.on_click(_toggle_edit_s4)
                        save_btn.on_click(_save_edit_s4)

            # Feedback + regenerate
            with ui.card().classes("dg-card-flat w-full"):
                ui.label("수정 요청").style("font-weight: 600; font-size: 14px; color: var(--dg-text-primary)")
                feedback_s4 = ui.textarea(
                    placeholder="예: KPI 목표를 좀 더 보수적으로, 운영 일정을 4주로 변경 등"
                ).classes("w-full dg-input").props("rows=2 outlined")

                with ui.row().classes("gap-3 items-center"):
                    regen_s4_btn = ui.button(
                        "재생성", icon="refresh",
                        on_click=lambda: _regenerate_proposal(feedback_s4.value),
                    ).classes("dg-btn-ghost")
                    regen_s4_spinner = ui.spinner(size="24px").classes("hidden")
                    regen_s4_status = ui.label("").classes("dg-progress-text hidden")

                _wizard_state["_regen_s4_spinner"] = regen_s4_spinner
                _wizard_state["_regen_s4_status"] = regen_s4_status
                _wizard_state["_regen_s4_btn"] = regen_s4_btn

            # Navigation
            _render_export_bar(
                4,
                docx_handler=lambda: _export_proposal_default(),
                pdf_handler=lambda: _export_proposal_default_pdf(),
            )

    async def _generate_proposal() -> None:
        pid = nicegui_app.storage.user.get("current_project_id")
        if not pid:
            ui.notify("프로젝트를 먼저 선택해 주세요.", type="warning")
            return
        project = get_project(pid)
        if not project:
            ui.notify("프로젝트를 찾을 수 없어요. 프로젝트 페이지에서 다시 선택해 주세요.", type="negative")
            return

        engine = engine_radio.value

        # Claude는 기본이 CLI(구독). API 백엔드일 때만 키 필요.
        if (engine == "claude" and os.getenv("CLAUDE_BACKEND", "cli").strip().lower() == "api"
                and not os.getenv("ANTHROPIC_API_KEY", "")):
            ui.notify("Claude를 API 모드로 쓰려면 .env에 ANTHROPIC_API_KEY가 필요해요. (CLI 모드는 키 없이 동작)", type="negative")
            return
        # GPT는 기본이 codex CLI(구독). API 백엔드일 때만 키 필요.
        if (engine in ("gpt", "coordinate") and os.getenv("OPENAI_BACKEND", "cli").strip().lower() == "api"
                and not os.getenv("OPENAI_API_KEY", "")):
            ui.notify("GPT를 API 모드로 쓰려면 .env에 OPENAI_API_KEY가 필요해요. (CLI 모드는 codex 로그인으로 동작)", type="negative")
            return

        btn = _wizard_state.get("_s4_btn")
        spinner = _wizard_state.get("_s4_spinner")
        status = _wizard_state.get("_s4_status")

        if btn:
            btn.props("disabled loading")
        if spinner:
            spinner.classes(remove="hidden")
        if status:
            status.classes(remove="hidden")
            status.set_text("운영 제안서를 만들고 있어요...")

        try:
            guide, prompt = build_wizard_proposal_prompt(
                project,
                strategy_context=_wizard_state.get("step1_content", ""),
                content_context=_wizard_state.get("step2_content", ""),
                ad_settings_context=_wizard_state.get("step3_content", ""),
            )
            loop = asyncio.get_running_loop()
            if engine == "coordinate":
                from app.ai.coordination import coordinate_generate
                content = await coordinate_generate(
                    loop, prompt, guide, "운영 제안서",
                    on_drafts=(lambda: status.set_text("Claude와 GPT가 각자 작성하고 있어요...")) if status else None,
                    on_synth=(lambda: status.set_text("Claude가 두 초안을 종합하고 있어요...")) if status else None,
                )
            else:
                provider = get_provider(engine)
                content = await loop.run_in_executor(
                    None, lambda: provider.generate_text(prompt, system_prompt=guide),
                )

            # 누락/부실 섹션이 있으면 1회 보정 (best-effort, 실패 시 원본 유지).
            if status:
                status.set_text("빠진 부분이 없는지 확인하고 있어요...")
            content = await loop.run_in_executor(
                None, lambda: repair_output(content, get_schema("wizard_proposal"), engine=engine),
            )

            _wizard_state["step4_content"] = content
            _wizard_state["step4_sections"] = parse_wizard_proposal_sections(content)
            save_generated_content(pid, engine, content, content_type="wizard_proposal")

            _render_proposal_result()
            ui.notify("운영 제안서가 완성됐어요!", type="positive")

        except Exception as exc:
            _log.exception("운영 제안서 생성 실패: %s", exc)
            ui.notify(f"운영 제안서를 만들지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative", timeout=8000)
            if status:
                status.set_text(f"오류: {exc}")
        finally:
            if btn:
                btn.props(remove="disabled loading")
            if spinner:
                spinner.classes("hidden")

    async def _regenerate_proposal(feedback: str = "") -> None:
        pid = nicegui_app.storage.user.get("current_project_id")
        if not pid:
            return
        project = get_project(pid)
        if not project:
            return

        engine = engine_radio.value if engine_radio.value not in ("both", "coordinate") else "claude"

        regen_spinner = _wizard_state.get("_regen_s4_spinner")
        regen_status = _wizard_state.get("_regen_s4_status")
        regen_btn = _wizard_state.get("_regen_s4_btn")

        if regen_btn:
            regen_btn.props("disabled loading")
        if regen_spinner:
            regen_spinner.classes(remove="hidden")
        if regen_status:
            regen_status.classes(remove="hidden")
            regen_status.set_text("다시 만들고 있어요...")

        try:
            guide, prompt = build_wizard_proposal_prompt(
                project,
                strategy_context=_wizard_state.get("step1_content", ""),
                content_context=_wizard_state.get("step2_content", ""),
                ad_settings_context=_wizard_state.get("step3_content", ""),
            )
            if feedback.strip():
                prompt += f"\n\n[수정 요청]\n{feedback.strip()}"
            if _wizard_state["step4_content"]:
                prompt += f"\n\n[이전 결과 (참고하되 수정 요청 반영)]\n{_wizard_state['step4_content']}"

            loop = asyncio.get_running_loop()
            provider = get_provider(engine)
            content = await loop.run_in_executor(
                None, lambda: provider.generate_text(prompt, system_prompt=guide),
            )

            _wizard_state["step4_content"] = content
            _wizard_state["step4_sections"] = parse_wizard_proposal_sections(content)
            save_generated_content(pid, engine, content, content_type="wizard_proposal")

            _render_proposal_result()
            ui.notify("운영 제안서를 다시 만들었어요!", type="positive")

        except Exception as exc:
            _log.exception("운영 제안서 재생성 실패: %s", exc)
            ui.notify(f"다시 만들지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative", timeout=8000)
        finally:
            if regen_btn:
                regen_btn.props(remove="disabled loading")
            if regen_spinner:
                regen_spinner.classes("hidden")

    # ── Generic step export handler ──────────────────────────────────────────

    async def _export_step_content(step_num: int, fmt: str = "docx") -> None:
        """범용 스텝 내보내기: 마크다운 내용을 DOCX 또는 PDF로 저장."""
        import tempfile
        from datetime import datetime
        from docx import Document

        pid = nicegui_app.storage.user.get("current_project_id")
        project = get_project(pid) if pid else None
        name = project.get("name", "unknown") if project else "unknown"

        content_key = f"step{step_num}_content"
        content = _wizard_state.get(content_key, "")
        if not content:
            ui.notify("콘텐츠를 먼저 만들어 주세요.", type="warning")
            return

        step_names = {1: "전략분석", 2: "콘텐츠", 3: "광고세팅", 4: "운영제안서"}
        step_name = step_names.get(step_num, f"Step{step_num}")

        try:
            doc = Document()
            doc.add_heading(f"{name} - {step_name}", level=1)
            doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph("")

            import re
            from docx.shared import Pt, Cm
            from docx.enum.table import WD_TABLE_ALIGNMENT

            def _add_md_paragraph(doc, text: str, style=None):
                """마크다운 서식(**bold**, *italic*)을 Word 서식으로 변환하여 추가."""
                # 마크다운 기호 제거 및 서식 변환
                p = doc.add_paragraph(style=style)
                # **bold** 와 *italic* 패턴을 분리
                parts = re.split(r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*)', text)
                for part in parts:
                    if part.startswith('***') and part.endswith('***'):
                        run = p.add_run(part[3:-3])
                        run.bold = True
                        run.italic = True
                    elif part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                        run = p.add_run(part[1:-1])
                        run.italic = True
                    else:
                        # 남은 마크다운 기호 정리
                        cleaned = part.replace('`', '').replace('~~', '')
                        if cleaned:
                            p.add_run(cleaned)
                return p

            def _clean_md(text: str) -> str:
                """마크다운 기호를 제거하고 순수 텍스트 반환."""
                text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                text = text.replace('`', '').replace('~~', '')
                return text

            lines = content.split('\n')
            i = 0
            while i < len(lines):
                stripped = lines[i].strip()

                # 마크다운 표 감지: | 로 시작하는 연속된 라인
                if stripped.startswith('|') and '|' in stripped[1:]:
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        line = lines[i].strip()
                        # 구분선(|---|---|) 건너뛰기
                        if not re.match(r'^\|[\s\-:]+\|', line):
                            table_lines.append(line)
                        i += 1

                    if table_lines:
                        headers = [_clean_md(cell.strip()) for cell in table_lines[0].split('|') if cell.strip()]
                        if headers:
                            table = doc.add_table(rows=1, cols=len(headers))
                            try:
                                table.style = 'Table Grid'
                            except KeyError:
                                pass  # fallback: no style
                            table.alignment = WD_TABLE_ALIGNMENT.CENTER

                            # 열 너비 균등 배분 (페이지 기준)
                            col_width = Cm(16 / len(headers)) if len(headers) > 0 else Cm(3)
                            for j, header in enumerate(headers):
                                cell = table.rows[0].cells[j]
                                cell.width = col_width
                                cell.text = header
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                                        run.font.size = Pt(9)

                            # 데이터 행
                            for row_line in table_lines[1:]:
                                cells = [_clean_md(cell.strip()) for cell in row_line.split('|') if cell.strip()]
                                row = table.add_row()
                                for j, cell_text in enumerate(cells):
                                    if j < len(headers):
                                        row.cells[j].width = col_width
                                        row.cells[j].text = cell_text
                                        for paragraph in row.cells[j].paragraphs:
                                            for run in paragraph.runs:
                                                run.font.size = Pt(9)

                            doc.add_paragraph("")
                    continue

                if stripped.startswith('### '):
                    doc.add_heading(_clean_md(stripped[4:]), level=3)
                elif stripped.startswith('## '):
                    doc.add_heading(_clean_md(stripped[3:]), level=2)
                elif stripped.startswith('- '):
                    _add_md_paragraph(doc, stripped[2:], style='List Bullet')
                elif stripped:
                    _add_md_paragraph(doc, stripped)

                i += 1

            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            fname = f"{step_name}_{name}_{ts}.docx"

            with tempfile.TemporaryDirectory() as tmpdir:
                tmp_path = Path(tmpdir) / fname
                doc.save(str(tmp_path))
                docx_bytes = tmp_path.read_bytes()

                if fmt == "pdf":
                    try:
                        from docx2pdf import convert
                        import pythoncom
                        pdf_fname = fname.replace(".docx", ".pdf")
                        pdf_path = Path(tmpdir) / pdf_fname
                        pythoncom.CoInitialize()
                        try:
                            convert(str(tmp_path), str(pdf_path))
                        finally:
                            pythoncom.CoUninitialize()
                        pdf_bytes = pdf_path.read_bytes()
                        ExportManager.save_default(pdf_bytes, pdf_fname)
                        ui.notify(f"{pdf_fname} 파일을 저장했어요.", type="positive", timeout=5000)
                    except ImportError:
                        ui.notify("PDF로 저장하려면 docx2pdf 설치가 필요해요. 터미널에서 pip install docx2pdf를 실행해 주세요.", type="warning")
                    except Exception as pdf_exc:
                        ui.notify(f"PDF로 바꾸지 못해서 DOCX로 저장했어요. ({pdf_exc})", type="warning")
                        ExportManager.save_default(docx_bytes, fname)
                        ui.notify(f"{fname} 파일을 저장했어요.", type="positive", timeout=5000)
                else:
                    ExportManager.save_default(docx_bytes, fname)
                    ui.notify(f"{fname} 파일을 저장했어요.", type="positive", timeout=5000)
        except Exception as exc:
            ui.notify(f"파일을 내보내지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative")

    # ── Proposal export handlers ─────────────────────────────────────────────

    def _build_proposal_docx_bytes(pid_snapshot: str | None = None) -> tuple:
        import tempfile
        from datetime import datetime as _dt
        content = _wizard_state.get("step4_content", "")
        if not content:
            raise ValueError("운영 제안서를 먼저 만들어 주세요.")
        project = get_project(pid_snapshot) if pid_snapshot else None
        if not project:
            raise ValueError("프로젝트를 선택해 주세요.")
        shop_info = {
            "shop_name": project.get("name", "광고주"),
            "industry": project.get("industry", ""),
            "location": project.get("region", ""),
        }
        ts = _dt.now().strftime("%Y%m%d_%H%M%S")
        filename = f"제안서_{shop_info['shop_name']}_{ts}.docx"
        sections = _wizard_state.get("step4_sections", {})
        # Map wizard proposal sections to proposal_docx expected keys
        docx_sections = {
            "summary": sections.get("summary", ""),
            "prev_performance": sections.get("target_summary", ""),
            "bottleneck": sections.get("content_strategy", ""),
            "kpi_goals": sections.get("budget_kpi", ""),
            "strategy": sections.get("execution_plan", ""),
            "execution": sections.get("schedule", ""),
            "creative": sections.get("judgment", ""),
        }
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir) / filename
            build_proposal_docx(
                shop_info=shop_info,
                sections=docx_sections,
                output_path=tmp_path,
            )
            docx_bytes = tmp_path.read_bytes()
        return docx_bytes, filename

    async def _export_proposal_default() -> None:
        export_btn = _wizard_state.get("_s4_export_default")
        dl_status = _wizard_state.get("_s4_download_status")
        if export_btn:
            export_btn.props("disabled loading")
        if dl_status:
            dl_status.classes(remove="hidden")
            dl_status.set_text("DOCX 파일을 준비하고 있어요...")
        try:
            pid_snap = nicegui_app.storage.user.get("current_project_id")
            loop = asyncio.get_running_loop()
            docx_bytes, fname = await loop.run_in_executor(None, lambda: _build_proposal_docx_bytes(pid_snap))
            ExportManager.save_default(docx_bytes, fname)
            if dl_status:
                dl_status.set_text(f"{fname} 저장 완료")
            ui.notify(f"{fname}", type="positive", timeout=8000, close_button="확인")
        except ValueError as ve:
            ui.notify(str(ve), type="warning")
        except Exception as exc:
            if dl_status:
                dl_status.set_text("내보내지 못했어요")
            ui.notify(f"파일을 내보내지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative")
        finally:
            if export_btn:
                export_btn.props(remove="disabled loading")

    async def _export_proposal_saveas() -> None:
        export_btn = _wizard_state.get("_s4_export_saveas")
        dl_status = _wizard_state.get("_s4_download_status")
        if export_btn:
            export_btn.props("disabled loading")
        if dl_status:
            dl_status.classes(remove="hidden")
            dl_status.set_text("DOCX 파일을 준비하고 있어요...")
        try:
            pid_snap = nicegui_app.storage.user.get("current_project_id")
            loop = asyncio.get_running_loop()
            docx_bytes, fname = await loop.run_in_executor(None, lambda: _build_proposal_docx_bytes(pid_snap))
            ok = await ExportManager.save_as(docx_bytes, fname)
            if dl_status:
                dl_status.set_text(f"{fname} 저장 완료" if ok else "저장 취소됨")
        except ValueError as ve:
            ui.notify(str(ve), type="warning")
        except Exception as exc:
            if dl_status:
                dl_status.set_text("내보내지 못했어요")
            ui.notify(f"파일을 내보내지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative")
        finally:
            if export_btn:
                export_btn.props(remove="disabled loading")

    async def _export_proposal_default_pdf() -> None:
        import tempfile
        export_btn = _wizard_state.get("_s4_export_pdf")
        dl_status = _wizard_state.get("_s4_download_status")
        if export_btn:
            export_btn.props("disabled loading")
        if dl_status:
            dl_status.classes(remove="hidden")
            dl_status.set_text("PDF 파일을 준비하고 있어요...")
        try:
            pid_snap = nicegui_app.storage.user.get("current_project_id")
            loop = asyncio.get_running_loop()
            docx_bytes, fname = await loop.run_in_executor(None, lambda: _build_proposal_docx_bytes(pid_snap))
            pdf_fname = fname.replace(".docx", ".pdf")

            def _convert_to_pdf():
                from docx2pdf import convert
                import pythoncom
                pythoncom.CoInitialize()
                try:
                    with tempfile.TemporaryDirectory() as tmpdir:
                        docx_path = Path(tmpdir) / fname
                        docx_path.write_bytes(docx_bytes)
                        pdf_path = Path(tmpdir) / pdf_fname
                        convert(str(docx_path), str(pdf_path))
                        return pdf_path.read_bytes()
                finally:
                    pythoncom.CoUninitialize()

            try:
                pdf_bytes = await loop.run_in_executor(None, _convert_to_pdf)
                ExportManager.save_default(pdf_bytes, pdf_fname)
                if dl_status:
                    dl_status.set_text(f"{pdf_fname} 저장 완료")
                ui.notify(f"{pdf_fname}", type="positive", timeout=8000, close_button="확인")
            except ImportError:
                ui.notify("PDF로 저장하려면 docx2pdf 설치가 필요해요. 터미널에서 pip install docx2pdf를 실행해 주세요.", type="warning")
                if dl_status:
                    dl_status.set_text("PDF 변환 불가 - docx2pdf 없음")
            except Exception as pdf_exc:
                ui.notify(f"PDF로 바꾸지 못해서 DOCX로 저장했어요. ({pdf_exc})", type="warning")
                ExportManager.save_default(docx_bytes, fname)
                if dl_status:
                    dl_status.set_text(f"{fname} 저장 완료 (DOCX 대체)")
        except ValueError as ve:
            ui.notify(str(ve), type="warning")
        except Exception as exc:
            if dl_status:
                dl_status.set_text("내보내지 못했어요")
            ui.notify(f"파일을 내보내지 못했어요. 잠시 후 다시 시도해 주세요. ({exc})", type="negative")
        finally:
            if export_btn:
                export_btn.props(remove="disabled loading")

    # ── Step router ─────────────────────────────────────────────────────────

    def _render_current_step() -> None:
        step = _wizard_state["current_step"]
        if step == 1:
            _render_step1()
        elif step == 2:
            _render_step2()
        elif step == 3:
            _render_step3()
        elif step == 4:
            _render_step4()

    # ── Initial render ──────────────────────────────────────────────────────

    # Load saved data from DB if available
    pid_init = nicegui_app.storage.user.get("current_project_id")
    if pid_init:
        saved_strategy = get_latest_content(pid_init, content_type="strategy")
        if saved_strategy:
            _wizard_state["step1_content"] = saved_strategy["content"]
            _wizard_state["step1_sections"] = parse_strategy_sections(saved_strategy["content"])
        saved_ad_settings = get_latest_content(pid_init, content_type="ad_settings")
        if saved_ad_settings:
            _wizard_state["step3_content"] = saved_ad_settings["content"]
            _wizard_state["step3_sections"] = parse_ad_settings_sections(saved_ad_settings["content"])
        saved_proposal = get_latest_content(pid_init, content_type="wizard_proposal")
        if saved_proposal:
            _wizard_state["step4_content"] = saved_proposal["content"]
            _wizard_state["step4_sections"] = parse_wizard_proposal_sections(saved_proposal["content"])

    _render_step_indicator()
    _render_current_step()

    # ── Public reset (called when project changes) ──────────────────────────

    def reset_wizard() -> None:
        """Reset wizard state for project switch."""
        _wizard_state["current_step"] = 1
        _wizard_state["step1_content"] = ""
        _wizard_state["step1_sections"] = {}
        _wizard_state["step2_content"] = ""
        _wizard_state["step2_engine"] = "claude"
        _wizard_state["step3_content"] = ""
        _wizard_state["step3_sections"] = {}
        _wizard_state["step4_content"] = ""
        _wizard_state["step4_sections"] = {}
        pid = nicegui_app.storage.user.get("current_project_id")
        _wizard_state["project_id"] = pid

        # Reload from DB
        if pid:
            saved_strategy = get_latest_content(pid, content_type="strategy")
            if saved_strategy:
                _wizard_state["step1_content"] = saved_strategy["content"]
                _wizard_state["step1_sections"] = parse_strategy_sections(saved_strategy["content"])
            saved_ad_settings = get_latest_content(pid, content_type="ad_settings")
            if saved_ad_settings:
                _wizard_state["step3_content"] = saved_ad_settings["content"]
                _wizard_state["step3_sections"] = parse_ad_settings_sections(saved_ad_settings["content"])
            saved_proposal = get_latest_content(pid, content_type="wizard_proposal")
            if saved_proposal:
                _wizard_state["step4_content"] = saved_proposal["content"]
                _wizard_state["step4_sections"] = parse_wizard_proposal_sections(saved_proposal["content"])

        _render_step_indicator()
        _render_current_step()

    # Return the reset function so planning.py can wire it up
    return reset_wizard
